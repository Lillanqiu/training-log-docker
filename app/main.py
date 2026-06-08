import asyncio
import base64
import hashlib
import hmac
import json
import os
import re
import secrets
import shutil
import subprocess
import time
import uuid
import zipfile
from datetime import datetime, timedelta
from io import BytesIO
from pathlib import Path
from typing import Any

import httpx
from openpyxl import load_workbook
from docx import Document
from docx.table import _Cell
from fastapi import FastAPI, File, Form, HTTPException, Request, Response, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pypdf import PdfReader


BASE_DIR = Path(__file__).resolve().parent
UPLOAD_DIR = Path("/app/uploads")
OUTPUT_DIR = Path("/app/outputs")
CONFIG_DIR = Path("/app/config")
CONFIG_PATH = CONFIG_DIR / "api-config.json"
USERS_PATH = CONFIG_DIR / "users.json"
AUTH_SETTINGS_PATH = CONFIG_DIR / "auth-settings.json"
DEFAULT_TEMPLATE_PATH = BASE_DIR / "templates_docx" / "default-training-log.docx"
MAX_TEXT_CHARS = 18000
MAX_REQUIREMENTS_CHARS = 10000
SESSION_COOKIE = "training_log_session"
SESSION_TTL_SECONDS = 60 * 60 * 12
JOBS: dict[str, dict[str, Any]] = {}
JOB_TASKS: dict[str, asyncio.Task[Any]] = {}

app = FastAPI(title="AI Document Form Filler")
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")


@app.exception_handler(Exception)
async def unhandled_exception_handler(_request: Request, exc: Exception) -> JSONResponse:
    return JSONResponse(
        status_code=500,
        content={"detail": f"Server error: {type(exc).__name__}: {exc}"},
    )


@app.get("/", response_class=HTMLResponse)
async def index() -> str:
    return (BASE_DIR / "templates" / "index.html").read_text(encoding="utf-8")


@app.get("/api/session")
async def get_session(request: Request) -> dict[str, Any]:
    if not is_auth_required():
        return {
            "authenticated": True,
            "auth_required": False,
            "user": public_user(guest_user()),
        }
    user = current_user_from_request(request)
    return {
        "authenticated": bool(user),
        "auth_required": True,
        "user": public_user(user) if user else None,
    }


@app.post("/api/login")
async def login(payload: dict[str, str], response: Response) -> dict[str, Any]:
    username = str(payload.get("username") or "").strip()
    password = str(payload.get("password") or "")
    user = authenticate_user(username, password)
    if not user:
        raise HTTPException(status_code=401, detail="用户名或密码不正确。")
    response.set_cookie(
        SESSION_COOKIE,
        create_session_token(user["username"]),
        max_age=SESSION_TTL_SECONDS,
        httponly=True,
        samesite="lax",
    )
    return {"authenticated": True, "user": public_user(user)}


@app.post("/api/logout")
async def logout(response: Response) -> dict[str, bool]:
    response.delete_cookie(SESSION_COOKIE)
    return {"logged_out": True}


@app.get("/api/auth-settings")
async def get_auth_settings(request: Request) -> dict[str, bool]:
    require_current_user(request)
    return {"require_password": is_auth_required()}


@app.post("/api/auth-settings")
async def save_auth_settings(payload: dict[str, bool], request: Request) -> dict[str, bool]:
    require_current_user(request)
    require_password = bool(payload.get("require_password", True))
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    AUTH_SETTINGS_PATH.write_text(json.dumps({"require_password": require_password}, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"require_password": require_password}


@app.get("/api/config")
async def get_config(request: Request) -> dict[str, str]:
    user = require_current_user(request)
    config_path = api_config_path_for_user(user["username"])
    if not config_path.exists():
        return {
            "format": "openai",
            "base_url": "https://api.openai.com/v1",
            "api_key": "",
            "model": "gpt-4.1-mini",
            "gpu_model": "",
        }
    return json.loads(config_path.read_text(encoding="utf-8"))


@app.post("/api/config")
async def save_config_endpoint(payload: dict[str, str], request: Request) -> dict[str, bool]:
    user = require_current_user(request)
    save_api_config(normalize_api_config(payload), user["username"])
    return {"saved": True}


@app.get("/api/account-backup")
async def export_account_backup(request: Request) -> dict[str, Any]:
    user = require_current_user(request)
    config_path = api_config_path_for_user(user["username"])
    api_config = json.loads(config_path.read_text(encoding="utf-8")) if config_path.exists() else {
        "format": "openai",
        "base_url": "https://api.openai.com/v1",
        "api_key": "",
        "model": "gpt-4.1-mini",
        "gpu_model": "",
    }
    default_template = None
    if DEFAULT_TEMPLATE_PATH.exists():
        default_template = {
            "filename": DEFAULT_TEMPLATE_PATH.name,
            "content_base64": base64.b64encode(DEFAULT_TEMPLATE_PATH.read_bytes()).decode("ascii"),
        }
    return {
        "version": 1,
        "exported_at": datetime.utcnow().isoformat(timespec="seconds") + "Z",
        "user": public_user(user),
        "server": {
            "api_config": api_config,
            "auth_settings": {"require_password": is_auth_required()},
            "default_template": default_template,
        },
    }


@app.post("/api/account-backup")
async def restore_account_backup(payload: dict[str, Any], request: Request) -> dict[str, bool]:
    user = require_current_user(request)
    server = payload.get("server") if isinstance(payload.get("server"), dict) else payload
    api_config = server.get("api_config") if isinstance(server, dict) else None
    if isinstance(api_config, dict):
        save_api_config(normalize_api_config(api_config), user["username"])
    auth_settings = server.get("auth_settings") if isinstance(server, dict) else None
    if isinstance(auth_settings, dict) and "require_password" in auth_settings:
        CONFIG_DIR.mkdir(parents=True, exist_ok=True)
        AUTH_SETTINGS_PATH.write_text(
            json.dumps({"require_password": bool(auth_settings.get("require_password"))}, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    default_template = server.get("default_template") if isinstance(server, dict) else None
    if isinstance(default_template, dict) and default_template.get("content_base64"):
        try:
            template_bytes = base64.b64decode(str(default_template["content_base64"]), validate=True)
        except Exception as exc:
            raise HTTPException(status_code=400, detail="默认填写模板不是有效的 base64 内容。") from exc
        if len(template_bytes) > 10 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="默认填写模板超过 10MB，已拒绝恢复。")
        DEFAULT_TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
        DEFAULT_TEMPLATE_PATH.write_bytes(template_bytes)
    return {"restored": True}


@app.post("/api/template-preview")
async def template_preview(request: Request, file: UploadFile | None = File(None)) -> dict[str, Any]:
    require_current_user(request)
    using_default = not (file and file.filename)
    if using_default:
        if not DEFAULT_TEMPLATE_PATH.exists():
            raise HTTPException(status_code=404, detail="默认训练日志模板不存在。")
        filename = "默认训练日志模板.docx"
        text = extract_text(DEFAULT_TEMPLATE_PATH, filename)
        fields = extract_docx_template_fields(DEFAULT_TEMPLATE_PATH) if DEFAULT_TEMPLATE_PATH.suffix.lower() == ".docx" else []
    else:
        filename = file.filename or "template"
        suffix = Path(filename).suffix.lower()
        if suffix not in {".pdf", ".docx", ".txt", ".md"}:
            raise HTTPException(status_code=400, detail="只支持预览 PDF、DOCX、TXT、MD 模板。")
        content = await file.read()
        if len(content) > 20 * 1024 * 1024:
            raise HTTPException(status_code=400, detail="模板文件不能超过 20MB。")
        text, fields = extract_template_preview_from_bytes(content, suffix)
    preview_text = text.strip()
    return {
        "filename": filename,
        "using_default": using_default,
        "fields": fields,
        "text_preview": preview_text[:5000],
        "truncated": len(preview_text) > 5000,
    }


@app.post("/api/test-api")
async def test_api_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    prompt = '请只返回 JSON：{"ok": true}'
    try:
        if config["format"] == "ollama":
            models = await list_ollama_models(config["base_url"], config["api_key"])
            if config["model"] not in models:
                raise HTTPException(
                    status_code=400,
                    detail=f"Ollama 已连接，但当前模型不存在：{config['model']}。可用模型：{', '.join(models[:20]) or '无'}",
                )
            return {
                "ok": True,
                "format": config["format"],
                "base_url": config["base_url"],
                "model": config["model"],
                "message": "Ollama 连接成功，当前模型可用。",
                "preview": f"已识别 {len(models)} 个本地模型。",
            }
        else:
            if not config["api_key"]:
                raise HTTPException(status_code=400, detail="OpenAI 兼容接口需要填写 API 密钥。")
            content = await call_openai_compatible_raw(config["base_url"], config["api_key"], config["model"], prompt)
        return {
            "ok": True,
            "format": config["format"],
            "base_url": config["base_url"],
            "model": config["model"],
            "message": "API 连接成功，模型有响应。",
            "preview": content[:300],
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"API 测试失败：{type(exc).__name__}: {exc}") from exc


@app.post("/api/preload-model")
async def preload_model_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    await ensure_ollama_model(config)
    await set_ollama_model_keep_alive(config["base_url"], config["model"], "30m", config["api_key"])
    return {
        "ok": True,
        "format": config["format"],
        "base_url": config["base_url"],
        "model": config["model"],
        "message": f"模型已预加载：{config['model']}。",
    }


@app.post("/api/restart-model")
async def restart_model_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    await ensure_ollama_model(config)
    await set_ollama_model_keep_alive(config["base_url"], config["model"], 0, config["api_key"])
    await set_ollama_model_keep_alive(config["base_url"], config["model"], "30m", config["api_key"])
    return {
        "ok": True,
        "format": config["format"],
        "base_url": config["base_url"],
        "model": config["model"],
        "message": f"模型已重启并重新加载：{config['model']}。",
    }


@app.post("/api/unload-model")
async def unload_model_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    await ensure_ollama_model(config)
    await set_ollama_model_keep_alive(config["base_url"], config["model"], 0, config["api_key"])
    return {
        "ok": True,
        "format": config["format"],
        "base_url": config["base_url"],
        "model": config["model"],
        "message": f"模型已关闭：{config['model']}。",
    }


@app.post("/api/fill-job")
async def create_fill_job(
    request: Request,
    file: UploadFile | None = File(None),
    form_schema: str = Form(""),
    requirements_text: str = Form(""),
    custom_prompt: str = Form(""),
    custom_skills: str = Form(""),
    start_date: str = Form(""),
    default_coach: str = Form(""),
    default_writer: str = Form(""),
    selected_modules: str = Form(""),
    link_module_teacher: bool = Form(False),
    module_teacher_map: str = Form(""),
    module_a_teacher: str = Form(""),
    module_b_teacher: str = Form(""),
    module_c_teacher: str = Form(""),
    selected_records: str = Form(""),
    output_format: str = Form("docx"),
    output_package_mode: str = Form("folder"),
    api_format: str = Form("openai"),
    api_base_url: str = Form(""),
    api_key: str = Form(""),
    api_model: str = Form(""),
) -> dict[str, Any]:
    require_current_user(request)
    saved_path, original_filename = await resolve_template_file(file)
    document_text = extract_text(saved_path, original_filename)
    fields = parse_form_schema(form_schema)
    if not fields and saved_path.suffix.lower() == ".docx":
        fields = extract_docx_template_fields(saved_path)
    if not fields:
        fields = default_training_log_fields()

    records = parse_selected_records(selected_records)
    if not records:
        raise HTTPException(status_code=400, detail="没有选中的训练对象，请先导入训练计划并勾选要生成的对象。")

    api_config = normalize_api_config(
        {
            "format": api_format,
            "base_url": api_base_url,
            "api_key": api_key,
            "model": api_model,
        }
    )
    defaults = {
        "start_date": start_date,
        "default_coach": default_coach,
        "default_writer": default_writer,
        "selected_modules": parse_selected_modules(selected_modules),
        "link_module_teacher": link_module_teacher,
        "module_teacher_map": build_module_teacher_map(module_teacher_map, module_a_teacher, module_b_teacher, module_c_teacher),
    }
    job_id = uuid.uuid4().hex
    JOBS[job_id] = {
        "id": job_id,
        "status": "running",
        "completed": 0,
        "total": len(records),
        "current": "",
        "message": "准备开始批量生成...",
        "cancel_requested": False,
        "result": None,
        "error": "",
    }
    JOB_TASKS[job_id] = asyncio.create_task(
        run_fill_job(
            job_id,
            saved_path,
            original_filename,
            document_text[:MAX_TEXT_CHARS],
            requirements_text[:MAX_REQUIREMENTS_CHARS],
            records,
            fields,
            api_config,
            defaults,
            custom_prompt,
            custom_skills,
            output_format,
            output_package_mode,
        )
    )
    return job_snapshot(job_id)


@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str, request: Request) -> dict[str, Any]:
    require_current_user(request)
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="任务不存在。")
    return job_snapshot(job_id)


@app.delete("/api/jobs/{job_id}")
async def cancel_job(job_id: str, request: Request) -> dict[str, Any]:
    require_current_user(request)
    if job_id not in JOBS:
        raise HTTPException(status_code=404, detail="任务不存在。")
    JOBS[job_id]["cancel_requested"] = True
    JOBS[job_id]["message"] = "正在取消，当前请求结束后停止..."
    task = JOB_TASKS.get(job_id)
    if task and not task.done():
        task.cancel()
    return job_snapshot(job_id)


@app.post("/api/models")
async def list_models_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    try:
        if config["format"] == "ollama":
            models = await list_ollama_models(config["base_url"], config["api_key"])
        else:
            if not config["api_key"]:
                raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
            models = await list_openai_compatible_models(config["base_url"], config["api_key"], config["model"])
        return {
            "ok": True,
            "format": config["format"],
            "base_url": config["base_url"],
            "models": models,
            "message": f"已获取 {len(models)} 个模型。",
        }
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"获取模型失败：{type(exc).__name__}: {exc}") from exc


@app.post("/api/gpus")
async def list_gpus_endpoint(payload: dict[str, str], request: Request) -> dict[str, Any]:
    require_current_user(request)
    config = normalize_api_config(payload)
    if config["format"] != "ollama":
        return {
            "ok": True,
            "format": config["format"],
            "gpus": [],
            "message": "GPU 读取只用于 Ollama 本地模型。",
        }
    gpus = await detect_nvidia_gpus()
    ollama_processors = await read_ollama_processors(config["base_url"])
    if gpus:
        message = "处理完成。"
        if ollama_processors:
            message += ""
    else:
        message = "处理完成。"
        if ollama_processors:
            message += ""
    return {
        "ok": True,
        "format": config["format"],
        "base_url": config["base_url"],
        "gpu_model": config.get("gpu_model", ""),
        "gpus": gpus,
        "ollama_processors": ollama_processors,
        "message": message,
    }


@app.get("/download/{filename}")
async def download_file(filename: str, request: Request) -> FileResponse:
    require_current_user(request)
    if not re.fullmatch(r"[a-f0-9]{32}(?:-[\w\u4e00-\u9fff.-]+)?\.(docx|pdf|zip)", filename):
        raise HTTPException(status_code=404, detail="File not found.")
    target = OUTPUT_DIR / filename
    if not target.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    if target.suffix == ".pdf":
        return FileResponse(
            target,
            media_type="application/pdf",
            filename=download_display_name(target),
        )
    if target.suffix == ".zip":
        return FileResponse(
            target,
            media_type="application/zip",
            filename=download_display_name(target),
        )
    return FileResponse(
        target,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=download_display_name(target),
    )


@app.get("/download-folder/{folder}/{filename}")
async def download_folder_file(folder: str, filename: str, request: Request) -> FileResponse:
    require_current_user(request)
    if not re.fullmatch(r"[a-f0-9]{32}(?:-[\w\u4e00-\u9fff.-]+)?", folder):
        raise HTTPException(status_code=404, detail="File not found.")
    if not re.fullmatch(r"[\w\u4e00-\u9fff .-]+\.(docx|pdf)", filename):
        raise HTTPException(status_code=404, detail="File not found.")
    folder_path = (OUTPUT_DIR / folder).resolve()
    output_root = OUTPUT_DIR.resolve()
    if output_root not in folder_path.parents and folder_path != output_root:
        raise HTTPException(status_code=404, detail="File not found.")
    target = folder_path / filename
    if not target.exists() or not target.is_file():
        raise HTTPException(status_code=404, detail="File not found.")
    if target.suffix == ".pdf":
        return FileResponse(target, media_type="application/pdf", filename=target.name)
    return FileResponse(
        target,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=target.name,
    )


@app.get("/download-folder-zip/{folder}")
async def download_folder_zip(folder: str, request: Request) -> FileResponse:
    require_current_user(request)
    folder_path = resolve_output_folder(folder)
    zip_path = zip_output_folder(folder_path)
    return FileResponse(
        zip_path,
        media_type="application/zip",
        filename=f"{safe_filename(folder_path.name)}.zip",
    )


@app.post("/api/fill")
async def fill_form(
    request: Request,
    file: UploadFile | None = File(None),
    form_schema: str = Form(""),
    requirements_url: str = Form(""),
    requirements_text: str = Form(""),
    custom_prompt: str = Form(""),
    custom_skills: str = Form(""),
    requirements_cookie: str = Form(""),
    requirements_user_agent: str = Form(""),
    requirements_referer: str = Form(""),
    start_date: str = Form(""),
    default_coach: str = Form(""),
    default_writer: str = Form(""),
    manual_module: str = Form(""),
    manual_training_content: str = Form(""),
    selected_modules: str = Form(""),
    link_module_teacher: bool = Form(False),
    module_teacher_map: str = Form(""),
    module_a_teacher: str = Form(""),
    module_b_teacher: str = Form(""),
    module_c_teacher: str = Form(""),
    selected_records: str = Form(""),
    batch_mode: bool = Form(False),
    save_config: bool = Form(False),
    output_format: str = Form("docx"),
    output_package_mode: str = Form("folder"),
    api_format: str = Form("openai"),
    api_base_url: str = Form(""),
    api_key: str = Form(""),
    api_model: str = Form(""),
) -> dict[str, Any]:
    user = require_current_user(request)
    saved_path, original_filename = await resolve_template_file(file)
    document_text = extract_text(saved_path, original_filename)
    if not document_text.strip():
        raise HTTPException(status_code=400, detail="No readable text was extracted from the document.")

    fields = parse_form_schema(form_schema)
    if not fields and saved_path.suffix.lower() == ".docx":
        fields = extract_docx_template_fields(saved_path)
    if not fields:
        raise HTTPException(status_code=400, detail="Please provide fields, or upload a DOCX template with detectable labels.")

    api_config = normalize_api_config(
        {
            "format": api_format,
            "base_url": api_base_url,
            "api_key": api_key,
            "model": api_model,
        }
    )
    if save_config:
        save_api_config(api_config, user["username"])

    requirements_text = requirements_text.strip() or await fetch_requirements_url(
        requirements_url,
        requirements_cookie,
        requirements_user_agent,
        requirements_referer,
    )
    if "kdocs.cn" in requirements_url.lower() and requirements_text.strip() and len(requirements_text.strip()) < 80:
        raise HTTPException(
            status_code=400,
            detail="金山文档链接没有返回可解析的表格正文。请在金山表格里复制“模块/训练计划”等内容，粘贴到“任务文档内容”框后再导入。",
        )
    imported_records = parse_selected_records(selected_records)
    apply_plan_defaults(
        imported_records,
        start_date,
        default_coach,
        default_writer,
        parse_selected_modules(selected_modules),
        link_module_teacher,
        build_module_teacher_map(module_teacher_map, module_a_teacher, module_b_teacher, module_c_teacher),
    )
    if batch_mode:
        if imported_records:
            if can_use_ai(api_config):
                batch_result = await generate_selected_records_one_by_one(
                    document_text[:MAX_TEXT_CHARS],
                    requirements_text[:MAX_REQUIREMENTS_CHARS],
                    imported_records,
                    fields,
                    api_config,
                    custom_prompt,
                    custom_skills,
                )
                apply_plan_defaults(
                    batch_result["records"],
                    start_date,
                    default_coach,
                    default_writer,
                    parse_selected_modules(selected_modules),
                    link_module_teacher,
                    build_module_teacher_map(module_teacher_map, module_a_teacher, module_b_teacher, module_c_teacher),
                )
            else:
                batch_result = {"used_ai": False, "records": imported_records}
                attach_stats_to_records(batch_result["records"], fallback_generation_stats(0))
        else:
            batch_result = await generate_batch_records(
                document_text[:MAX_TEXT_CHARS],
                requirements_text[:MAX_REQUIREMENTS_CHARS],
                fields,
                api_config,
                custom_prompt,
                custom_skills,
            )
        for record in batch_result["records"]:
            record["answers"] = enrich_answers_for_fields(record.get("answers", []), fields)
        download_url = ""
        folder_path = ""
        if saved_path.suffix.lower() == ".docx":
            output_path = package_batch_outputs(saved_path, batch_result["records"], output_format, output_package_mode)
            if output_path.is_dir():
                folder_path = output_folder_display_path(output_path)
                folder_files = output_folder_file_links(output_path)
                folder_zip_url = output_folder_zip_url(output_path)
            else:
                download_url = f"/download/{output_path.name}"
                folder_files = []
                folder_zip_url = ""
        else:
            folder_files = []
            folder_zip_url = ""
        return {
            "filename": original_filename,
            "used_ai": batch_result["used_ai"],
            "batch": True,
            "records": batch_result["records"],
            "answers": batch_result["records"][0]["answers"] if batch_result["records"] else [],
            "download_url": download_url,
            "folder_path": folder_path,
            "folder_files": folder_files,
            "folder_zip_url": folder_zip_url,
            "source_preview": document_text[:1200],
            "requirements_preview": requirements_text[:1200],
        }

    if manual_module.strip() or manual_training_content.strip() or start_date.strip() or default_coach.strip() or default_writer.strip():
        manual_plan = manual_training_content.strip() or manual_module.strip() or "手动填写训练内容"
        manual_record = build_training_record(manual_module.strip(), manual_plan, fields, 1)
        if start_date.strip():
            set_record_answer(manual_record, "日期", start_date.strip(), 0.9, "来自手动填写。")
        if default_coach.strip():
            set_record_answer(manual_record, "授课老师", default_coach.strip(), 0.9, "来自手动填写。")
            set_record_answer(manual_record, "负责教练", default_coach.strip(), 0.9, "来自手动填写。")
        if default_writer.strip():
            set_record_answer(manual_record, "填写人", default_writer.strip(), 0.9, "来自手动填写。")
        if manual_module.strip():
            set_record_answer(manual_record, "模块", manual_module.strip(), 0.9, "来自手动填写。")
            set_record_answer(manual_record, "训练模块", manual_module.strip(), 0.9, "来自手动填写。")
        result = {"used_ai": False, "answers": manual_record["answers"]}
        result["generation_stats"] = fallback_generation_stats(0)
    else:
        result = await generate_answers(
            document_text[:MAX_TEXT_CHARS],
            requirements_text[:MAX_REQUIREMENTS_CHARS],
            fields,
            api_config,
            custom_prompt,
            custom_skills,
        )
    result["answers"] = enrich_answers_for_fields(result.get("answers", []), fields)
    download_url = ""
    if saved_path.suffix.lower() == ".docx":
        output_path = fill_docx_template(saved_path, result["answers"], title_from_answers(result["answers"], original_filename))
        output_path = convert_output_format(output_path, output_format)
        download_url = f"/download/{output_path.name}"

    return {
        "filename": original_filename,
        "used_ai": result["used_ai"],
        "batch": False,
        "answers": result["answers"],
        "generation_stats": result.get("generation_stats", fallback_generation_stats(0)),
        "download_url": download_url,
        "source_preview": document_text[:1200],
        "requirements_preview": requirements_text[:1200],
    }


@app.post("/api/import-plan")
async def import_plan(
    request: Request,
    file: UploadFile | None = File(None),
    plan_file: UploadFile | None = File(None),
    form_schema: str = Form(""),
    requirements_url: str = Form(""),
    requirements_text: str = Form(""),
    custom_prompt: str = Form(""),
    custom_skills: str = Form(""),
    requirements_cookie: str = Form(""),
    requirements_user_agent: str = Form(""),
    requirements_referer: str = Form(""),
    start_date: str = Form(""),
    default_coach: str = Form(""),
    default_writer: str = Form(""),
    selected_modules: str = Form(""),
    link_module_teacher: bool = Form(False),
    module_teacher_map: str = Form(""),
    module_a_teacher: str = Form(""),
    module_b_teacher: str = Form(""),
    module_c_teacher: str = Form(""),
    save_config: bool = Form(False),
    api_format: str = Form("openai"),
    api_base_url: str = Form(""),
    api_key: str = Form(""),
    api_model: str = Form(""),
) -> dict[str, Any]:
    user = require_current_user(request)
    api_config = normalize_api_config(
        {
            "format": api_format,
            "base_url": api_base_url,
            "api_key": api_key,
            "model": api_model,
        }
    )
    if save_config:
        save_api_config(api_config, user["username"])

    fields = parse_form_schema(form_schema)
    template_text = ""
    if file and file.filename:
        saved_path = await save_upload(file)
        template_text = extract_text(saved_path, file.filename)
        if not fields and saved_path.suffix.lower() == ".docx":
            fields = extract_docx_template_fields(saved_path)

    if not fields:
        fields = default_training_log_fields()

    plan_records = await parse_plan_upload(plan_file, fields)
    if plan_records:
        apply_plan_defaults(
            plan_records,
            start_date,
            default_coach,
            default_writer,
            parse_selected_modules(selected_modules),
            link_module_teacher,
            build_module_teacher_map(module_teacher_map, module_a_teacher, module_b_teacher, module_c_teacher),
        )
        return {
            "used_ai": False,
            "records": plan_records,
            "requirements_preview": f"从上传的训练计划文件解析出 {len(plan_records)} 条记录。",
        }

    requirements_text = requirements_text.strip() or await fetch_requirements_url(
        requirements_url,
        requirements_cookie,
        requirements_user_agent,
        requirements_referer,
    )
    if len(requirements_text) < 30:
        raise HTTPException(
            status_code=400,
            detail="链接内容太少，可能需要登录或没有公开访问权限。请确认链接可公开查看，或把任务文档导出后上传。",
        )

    result = await generate_batch_records(
        template_text[:MAX_TEXT_CHARS],
        requirements_text[:MAX_REQUIREMENTS_CHARS],
        fields,
        api_config,
        custom_prompt,
        custom_skills,
    )
    apply_plan_defaults(
        result["records"],
        start_date,
        default_coach,
        default_writer,
        parse_selected_modules(selected_modules),
        link_module_teacher,
        build_module_teacher_map(module_teacher_map, module_a_teacher, module_b_teacher, module_c_teacher),
    )
    return {
        "used_ai": result["used_ai"],
        "records": result["records"],
        "requirements_preview": requirements_text[:1200],
    }


def normalize_api_config(raw: dict[str, str]) -> dict[str, str]:
    api_format = (raw.get("format") or "openai").strip()
    if api_format not in {"openai", "ollama"}:
        raise HTTPException(status_code=400, detail="api_format only supports openai or ollama.")
    default_base = "http://host.docker.internal:11434" if api_format == "ollama" else "https://api.openai.com/v1"
    default_model = "qwen2.5:7b" if api_format == "ollama" else "gpt-4.1-mini"
    model = (raw.get("model") or os.getenv("AI_MODEL") or default_model).strip()
    if api_format == "ollama":
        model = normalize_ollama_model_name(model)
    return {
        "format": api_format,
        "base_url": normalize_docker_host_url((raw.get("base_url") or os.getenv("AI_BASE_URL") or default_base).strip().rstrip("/")),
        "api_key": (raw.get("api_key") or os.getenv("AI_API_KEY") or "").strip(),
        "model": model,
        "gpu_model": (raw.get("gpu_model") or "").strip(),
    }


def require_current_user(request: Request) -> dict[str, str]:
    if not is_auth_required():
        return guest_user()
    user = current_user_from_request(request)
    if not user:
        raise HTTPException(status_code=401, detail="请求失败，请检查输入或 API 配置。")
    return user


def is_auth_required() -> bool:
    if not AUTH_SETTINGS_PATH.exists():
        return False
    try:
        payload = json.loads(AUTH_SETTINGS_PATH.read_text(encoding="utf-8"))
        return bool(payload.get("require_password", False))
    except (json.JSONDecodeError, OSError):
        return False


def guest_user() -> dict[str, str]:
    return {"username": "guest", "display_name": "免登录用户", "password": "", "password_hash": ""}


def current_user_from_request(request: Request) -> dict[str, str] | None:
    token = request.cookies.get(SESSION_COOKIE, "")
    username = verify_session_token(token)
    if not username:
        return None
    return find_user(username)


def public_user(user: dict[str, str] | None) -> dict[str, str] | None:
    if not user:
        return None
    return {
        "username": user["username"],
        "display_name": user.get("display_name") or user["username"],
    }


def authenticate_user(username: str, password: str) -> dict[str, str] | None:
    user = find_user(username)
    if not user:
        return None
    expected_hash = user.get("password_hash", "")
    if expected_hash:
        actual_hash = hashlib.sha256(password.encode("utf-8")).hexdigest()
        return user if hmac.compare_digest(expected_hash, actual_hash) else None
    expected_password = user.get("password", "")
    return user if expected_password and hmac.compare_digest(expected_password, password) else None


def find_user(username: str) -> dict[str, str] | None:
    normalized = username.strip().lower()
    return next((user for user in load_users() if user["username"].lower() == normalized), None)


def load_users() -> list[dict[str, str]]:
    users = []
    if USERS_PATH.exists():
        try:
            raw_users = json.loads(USERS_PATH.read_text(encoding="utf-8"))
            if isinstance(raw_users, dict):
                raw_users = raw_users.get("users", [])
            if isinstance(raw_users, list):
                users = raw_users
        except (json.JSONDecodeError, OSError):
            users = []
    if not users:
        default_admin_hash = os.getenv("DEFAULT_ADMIN_PASSWORD_HASH", "").strip()
        default_admin_password = os.getenv("DEFAULT_ADMIN_PASSWORD", "").strip()
        if default_admin_hash or default_admin_password:
            users = [{
                "username": os.getenv("DEFAULT_ADMIN_USERNAME", "admin").strip() or "admin",
                "display_name": os.getenv("DEFAULT_ADMIN_DISPLAY_NAME", "管理员").strip() or "管理员",
                "password": default_admin_password,
                "password_hash": default_admin_hash,
            }]
    normalized_users = []
    for user in users:
        if not isinstance(user, dict):
            continue
        username = str(user.get("username") or "").strip()
        if not username:
            continue
        normalized_users.append(
            {
                "username": username,
                "display_name": str(user.get("display_name") or username).strip(),
                "password": str(user.get("password") or ""),
                "password_hash": str(user.get("password_hash") or ""),
            }
        )
    return normalized_users


def create_session_token(username: str) -> str:
    expires_at = str(int(datetime.utcnow().timestamp()) + SESSION_TTL_SECONDS)
    payload = f"{username}.{expires_at}"
    signature = sign_text(payload)
    return base64.urlsafe_b64encode(f"{payload}.{signature}".encode("utf-8")).decode("ascii")


def verify_session_token(token: str) -> str:
    if not token:
        return ""
    try:
        decoded = base64.urlsafe_b64decode(token.encode("ascii")).decode("utf-8")
        username, expires_at, signature = decoded.rsplit(".", 2)
        payload = f"{username}.{expires_at}"
        if not hmac.compare_digest(signature, sign_text(payload)):
            return ""
        if int(expires_at) < int(datetime.utcnow().timestamp()):
            return ""
        return username
    except (ValueError, TypeError, OSError):
        return ""


def sign_text(value: str) -> str:
    return hmac.new(session_secret().encode("utf-8"), value.encode("utf-8"), hashlib.sha256).hexdigest()


def session_secret() -> str:
    env_secret = os.getenv("SESSION_SECRET", "").strip()
    if env_secret:
        return env_secret
    secret_path = CONFIG_DIR / "session-secret.txt"
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    if secret_path.exists():
        return secret_path.read_text(encoding="utf-8").strip()
    secret = secrets.token_urlsafe(32)
    secret_path.write_text(secret, encoding="utf-8")
    return secret


def api_config_path_for_user(username: str) -> Path:
    return CONFIG_DIR / f"api-config-{safe_config_name(username)}.json"


def safe_config_name(value: str) -> str:
    safe = re.sub(r"[^A-Za-z0-9_.-]+", "-", value.strip()).strip(".-")
    return safe or "user"


def normalize_docker_host_url(url: str) -> str:
    return re.sub(r"^(https?://)(localhost|127\.0\.0\.1)(:\d+)?", r"\1host.docker.internal\3", url, flags=re.IGNORECASE)


def normalize_ollama_model_name(model: str) -> str:
    model = model.strip()
    if ":" in model or not model:
        return model
    for suffix in ("latest", "128b", "72b", "70b", "32b", "31b", "27b", "26b", "14b", "13b", "8b", "7b", "4b", "3b", "1.5b"):
        if model.lower().endswith(suffix) and len(model) > len(suffix):
            return f"{model[:-len(suffix)]}:{model[-len(suffix):]}"
    return model


def can_use_ai(api_config: dict[str, str]) -> bool:
    if api_config["format"] == "ollama":
        return True
    return bool(api_config["api_key"])


def estimate_token_count(text: str) -> int:
    text = str(text or "")
    if not text.strip():
        return 0
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    ascii_words = len(re.findall(r"[A-Za-z0-9_]+", text))
    other_chars = max(0, len(text) - chinese_chars)
    return max(1, int(chinese_chars * 1.15 + ascii_words * 1.25 + other_chars / 4))


def build_generation_stats(
    provider: str,
    model: str,
    prompt: str = "",
    content: str = "",
    usage: dict[str, Any] | None = None,
    duration_seconds: float = 0,
    token_source: str = "estimated",
) -> dict[str, Any]:
    usage = usage or {}
    input_tokens = int(
        usage.get("prompt_tokens")
        or usage.get("input_tokens")
        or usage.get("prompt_eval_count")
        or usage.get("prompt_tokens_estimated")
        or 0
    )
    output_tokens = int(
        usage.get("completion_tokens")
        or usage.get("output_tokens")
        or usage.get("eval_count")
        or usage.get("completion_tokens_estimated")
        or 0
    )
    if not input_tokens:
        input_tokens = estimate_token_count(prompt)
        token_source = "estimated" if token_source == "api" else token_source
    if not output_tokens:
        output_tokens = estimate_token_count(content)
        token_source = "estimated" if token_source == "api" else token_source
    total_tokens = int(usage.get("total_tokens") or usage.get("total_tokens_estimated") or input_tokens + output_tokens)
    return {
        "provider": provider,
        "model": model,
        "duration_seconds": round(float(duration_seconds or 0), 2),
        "input_tokens": input_tokens,
        "output_tokens": output_tokens,
        "total_tokens": total_tokens,
        "token_source": token_source,
    }


def fallback_generation_stats(duration_seconds: float = 0) -> dict[str, Any]:
    return {
        "provider": "local-fallback",
        "model": "本地兜底",
        "duration_seconds": round(float(duration_seconds or 0), 2),
        "input_tokens": 0,
        "output_tokens": 0,
        "total_tokens": 0,
        "token_source": "none",
    }


def scale_generation_stats(stats: dict[str, Any], count: int) -> list[dict[str, Any]]:
    if count <= 0:
        return []
    input_each = round((int(stats.get("input_tokens") or 0)) / count)
    output_each = round((int(stats.get("output_tokens") or 0)) / count)
    total_each = round((int(stats.get("total_tokens") or 0)) / count)
    duration_each = round((float(stats.get("duration_seconds") or 0)) / count, 2)
    return [
        {
            **stats,
            "duration_seconds": duration_each,
            "input_tokens": input_each,
            "output_tokens": output_each,
            "total_tokens": total_each,
            "token_source": "allocated" if stats.get("token_source") != "none" else "none",
        }
        for _ in range(count)
    ]


def attach_stats_to_records(records: list[dict[str, Any]], stats: dict[str, Any] | list[dict[str, Any]] | None) -> None:
    if not records or not stats:
        return
    stats_list = stats if isinstance(stats, list) else scale_generation_stats(stats, len(records))
    for record, record_stats in zip(records, stats_list):
        record["generation_stats"] = record_stats


def build_selected_records_requirements(records: list[dict[str, Any]], requirements_text: str) -> str:
    return json.dumps(
        {
            "instruction": "请根据输入生成训练日志内容。",
            "selected_records": records,
            "source_requirements_text": requirements_text,
        },
        ensure_ascii=False,
    )


def job_snapshot(job_id: str) -> dict[str, Any]:
    job = JOBS[job_id]
    return {key: value for key, value in job.items() if key != "cancel_requested"}


async def run_fill_job(
    job_id: str,
    template_path: Path,
    original_filename: str,
    document_text: str,
    requirements_text: str,
    selected_records: list[dict[str, Any]],
    fields: list[dict[str, Any]],
    api_config: dict[str, str],
    defaults: dict[str, Any],
    custom_prompt: str,
    custom_skills: str,
    output_format: str,
    output_package_mode: str,
) -> None:
    job = JOBS[job_id]
    generated_paths: list[tuple[Path, str]] = []
    generated_records: list[dict[str, Any]] = []
    used_names: set[str] = set()
    used_ai = can_use_ai(api_config)
    try:
        for index, selected_record in enumerate(selected_records, start=1):
            if job.get("cancel_requested"):
                raise asyncio.CancelledError()
            title = str(selected_record.get("title") or f"training-log-{index}")
            job["current"] = title
            job["message"] = f"正在生成第 {index}/{len(selected_records)} 篇：{title}"
            if used_ai:
                result = await generate_selected_records_one_by_one(
                    document_text,
                    requirements_text,
                    [selected_record],
                    fields,
                    api_config,
                    custom_prompt,
                    custom_skills,
                )
                record = result["records"][0] if result["records"] else selected_record
            else:
                record = selected_record
                record["generation_stats"] = fallback_generation_stats(0)
            apply_plan_defaults(
                [record],
                defaults["start_date"],
                defaults["default_coach"],
                defaults["default_writer"],
                defaults["selected_modules"],
                defaults["link_module_teacher"],
                defaults["module_teacher_map"],
            )
            record["answers"] = enrich_answers_for_fields(record.get("answers", []), fields)
            output_title = record_file_title(record, index)
            output_path = fill_docx_template(template_path, record.get("answers", []), output_title)
            output_path = convert_output_format(output_path, output_format)
            file_name = unique_package_filename(f"{safe_filename(output_title)}{output_path.suffix}", used_names)
            generated_paths.append((output_path, file_name))
            generated_records.append(record)
            job["completed"] = index
            job["message"] = f"已写入第 {index}/{len(selected_records)} 篇：{title}"

        job["message"] = "正在整理下载文件..."
        package_path = write_generated_package(generated_paths, generated_records, output_package_mode)
        job["status"] = "complete"
        job["message"] = "批量生成完成。"
        job["result"] = {
            "filename": original_filename,
            "used_ai": used_ai,
            "batch": True,
            "records": generated_records,
            "answers": generated_records[0]["answers"] if generated_records else [],
            "download_url": "" if package_path.is_dir() else f"/download/{package_path.name}",
            "folder_path": output_folder_display_path(package_path) if package_path.is_dir() else "",
            "folder_files": output_folder_file_links(package_path) if package_path.is_dir() else [],
            "folder_zip_url": output_folder_zip_url(package_path) if package_path.is_dir() else "",
            "source_preview": document_text[:1200],
            "requirements_preview": requirements_text[:1200],
        }
    except asyncio.CancelledError:
        job["status"] = "canceled"
        job["message"] = "已取消批量生成。"
    except HTTPException as exc:
        job["status"] = "failed"
        job["error"] = str(exc.detail)
        job["message"] = str(exc.detail)
        attach_partial_batch_result(
            job,
            generated_paths,
            generated_records,
            output_package_mode,
            original_filename,
            used_ai,
            document_text,
            requirements_text,
        )
    except Exception as exc:
        job["status"] = "failed"
        job["error"] = f"{type(exc).__name__}: {exc}"
        job["message"] = str(exc)
        attach_partial_batch_result(
            job,
            generated_paths,
            generated_records,
            output_package_mode,
            original_filename,
            used_ai,
            document_text,
            requirements_text,
        )


def attach_partial_batch_result(
    job: dict[str, Any],
    generated_paths: list[tuple[Path, str]],
    generated_records: list[dict[str, Any]],
    output_package_mode: str,
    original_filename: str,
    used_ai: bool,
    document_text: str,
    requirements_text: str,
) -> None:
    if not generated_paths or not generated_records:
        return
    package_path = write_generated_package(generated_paths, generated_records, output_package_mode)
    job["partial_result"] = {
        "filename": original_filename,
        "used_ai": used_ai,
        "batch": True,
        "partial": True,
        "records": generated_records,
        "answers": generated_records[0]["answers"] if generated_records else [],
        "download_url": "" if package_path.is_dir() else f"/download/{package_path.name}",
        "folder_path": output_folder_display_path(package_path) if package_path.is_dir() else "",
        "folder_files": output_folder_file_links(package_path) if package_path.is_dir() else [],
        "folder_zip_url": output_folder_zip_url(package_path) if package_path.is_dir() else "",
        "source_preview": document_text[:1200],
        "requirements_preview": requirements_text[:1200],
    }
    job["message"] = f"{job.get('message') or '批量生成失败'} 已导出 {len(generated_records)} 个已完成文件。"


def write_generated_package(generated_paths: list[tuple[Path, str]], records: list[dict[str, Any]], package_mode: str) -> Path:
    root_folder = safe_filename(package_title_for_records(records))
    if normalize_package_mode(package_mode) == "folder":
        folder_path = unique_output_folder(root_folder)
        folder_path.mkdir(parents=True, exist_ok=False)
        for path, file_name in generated_paths:
            shutil.copy2(path, folder_path / file_name)
        return folder_path

    zip_path = OUTPUT_DIR / f"{uuid.uuid4().hex}-{root_folder}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path, file_name in generated_paths:
            arcname = f"{root_folder}/{file_name}"
            archive.write(path, arcname=arcname)
    return zip_path


def normalize_package_mode(package_mode: str) -> str:
    mode = (package_mode or "folder").strip().lower()
    if mode in {"zip", "zip_folder", "archive", "archive_folder"}:
        return "zip_folder"
    return "folder"


def unique_output_folder(root_folder: str) -> Path:
    base_name = safe_filename(root_folder) or "training-logs"
    candidate = OUTPUT_DIR / f"{uuid.uuid4().hex}-{base_name}"
    counter = 2
    while candidate.exists():
        candidate = OUTPUT_DIR / f"{uuid.uuid4().hex}-{base_name}-{counter}"
        counter += 1
    return candidate


def output_folder_display_path(path: Path) -> str:
    return f"outputs/{path.name}"


def output_folder_zip_url(path: Path) -> str:
    return f"/download-folder-zip/{path.name}" if path.is_dir() else ""


def resolve_output_folder(folder: str) -> Path:
    if not re.fullmatch(r"[a-f0-9]{32}(?:-[\w\u4e00-\u9fff.-]+)?", folder):
        raise HTTPException(status_code=404, detail="File not found.")
    folder_path = (OUTPUT_DIR / folder).resolve()
    output_root = OUTPUT_DIR.resolve()
    if output_root not in folder_path.parents and folder_path != output_root:
        raise HTTPException(status_code=404, detail="File not found.")
    if not folder_path.exists() or not folder_path.is_dir():
        raise HTTPException(status_code=404, detail="File not found.")
    return folder_path


def zip_output_folder(folder_path: Path) -> Path:
    files = [
        child for child in sorted(folder_path.iterdir(), key=lambda item: item.name)
        if child.is_file() and child.suffix.lower() in {".docx", ".pdf"}
    ]
    if not files:
        raise HTTPException(status_code=404, detail="File not found.")
    zip_path = OUTPUT_DIR / f"{uuid.uuid4().hex}-{safe_filename(folder_path.name)}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for child in files:
            archive.write(child, arcname=f"{folder_path.name}/{child.name}")
    return zip_path


def output_folder_file_links(path: Path) -> list[dict[str, str]]:
    if not path.is_dir():
        return []
    files = []
    for child in sorted(path.iterdir(), key=lambda item: item.name):
        if child.is_file() and child.suffix.lower() in {".docx", ".pdf"}:
            files.append({
                "name": child.name,
                "download_url": f"/download-folder/{path.name}/{child.name}",
            })
    return files


def default_training_log_fields() -> list[dict[str, Any]]:
    names = ["日期", "星期", "时间", "模块", "训练计划", "负责教练", "授课老师", "训练模块", "训练任务", "训练内容", "学习笔记", "总结", "填写人"]
    return [{"name": name, "options": []} for name in names]


async def parse_plan_upload(plan_file: UploadFile | None, fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not plan_file or not plan_file.filename:
        return []
    suffix = Path(plan_file.filename).suffix.lower()
    content = await plan_file.read()
    if not content:
        return []
    if suffix in {".xlsx", ".xlsm"}:
        return parse_plan_xlsx(content, fields)
    if suffix in {".txt", ".csv", ".tsv"}:
        return parse_training_plan_table(content.decode("utf-8", errors="ignore"), fields)
    raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")


def parse_plan_xlsx(content: bytes, fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    workbook = load_workbook(BytesIO(content), data_only=True, read_only=True)
    sheet = workbook.active
    rows = [
        [clean_excel_value(cell) for cell in row]
        for row in sheet.iter_rows(values_only=True)
    ]
    rows = [[cell for cell in row] for row in rows if any(row)]
    if not rows:
        return []

    header_index = find_plan_header_index(rows)
    if header_index is None:
        return parse_training_plan_table("\n".join("\t".join(row) for row in rows), fields)

    headers = rows[header_index]
    module_index = find_column(headers, ["模块", "训练模块", "分类", "项目", "方向", "类别"])
    plan_index = find_column(headers, ["训练计划", "训练任务", "训练内容", "任务", "计划", "内容", "每日任务", "日常训练计划"])
    date_index = find_column(headers, ["日期", "训练日期", "时间", "日程"])
    weekday_index = find_column(headers, ["星期", "周几"])
    time_index = find_column(headers, ["具体时间", "训练时间", "时段"])
    teacher_index = find_column(headers, ["负责教练", "授课老师", "老师", "教练", "负责人"])
    if plan_index is None:
        return []

    row_items = []
    last_date = ""
    last_weekday = ""
    last_module = ""
    last_teacher = ""
    for row in rows[header_index + 1:]:
        if plan_index >= len(row):
            continue
        plan = row[plan_index].strip()
        if not plan or clean_cell_text(plan) in {"训练计划", "训练任务", "训练内容"}:
            continue
        module = row[module_index].strip() if module_index is not None and module_index < len(row) else ""
        date_value = row[date_index].strip() if date_index is not None and date_index < len(row) else ""
        weekday_value = row[weekday_index].strip() if weekday_index is not None and weekday_index < len(row) else ""
        teacher_value = row[teacher_index].strip() if teacher_index is not None and teacher_index < len(row) else ""
        if date_value:
            last_date = date_value
        if weekday_value:
            last_weekday = weekday_value
        if module:
            last_module = module
        if teacher_value:
            last_teacher = teacher_value
        row_items.append({
            "module": module or last_module,
            "plan": plan,
            "date": date_value or last_date,
            "weekday": weekday_value or last_weekday,
            "time": row[time_index] if time_index is not None and time_index < len(row) else "",
            "teacher": teacher_value or last_teacher,
            "source": {
                "row": header_index + len(row_items) + 2,
                "module_column": headers[module_index] if module_index is not None and module_index < len(headers) else "",
                "plan_column": headers[plan_index] if plan_index < len(headers) else "",
                "teacher_column": headers[teacher_index] if teacher_index is not None and teacher_index < len(headers) else "",
            },
        })
    return group_plan_items_by_day(row_items, fields)[:31]


def group_plan_items_by_day(items: list[dict[str, Any]], fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not items:
        return []
    grouped: dict[str, list[dict[str, Any]]] = {}
    for index, item in enumerate(items):
        day_key = normalize_day_key(item.get("date", ""))
        key = day_key or f"row-group-{index // 3 + 1}"
        grouped.setdefault(key, []).append(item)

    records = []
    for key, group in grouped.items():
        modules = dedupe([str(item.get("module", "")) for item in group if item.get("module")])
        plans = dedupe([str(item.get("plan", "")).strip() for item in group if item.get("plan")])
        first = group[0]
        module = " / ".join(modules)
        plan_summary = "\n".join(f"{index}. {plan}" for index, plan in enumerate(plans, start=1))
        record = build_training_record(module, plan_summary, fields, len(records) + 1)
        record["date"] = str(first.get("date", ""))
        apply_optional_answer(record, "日期", str(first.get("date", "")))
        apply_optional_answer(record, "星期", str(first.get("weekday", "")))
        apply_optional_answer(record, "时间", str(first.get("time", "")))
        apply_optional_answer(record, "授课老师", str(first.get("teacher", "")))
        apply_optional_answer(record, "负责教练", str(first.get("teacher", "")))
        record["title"] = first_plan_title(plans) or f"{len(records) + 1:02d}-{first.get('date') or '训练日'}-{modules[0] if modules else '训练'}"
        record["source"] = {
            "row": f"{first.get('source', {}).get('row', '-')}-{group[-1].get('source', {}).get('row', '-')}",
            "plan_column": first.get("source", {}).get("plan_column", ""),
            "grouped_rows": len(group),
        }
        records.append(record)
    return records


def normalize_day_key(value: str) -> str:
    value = str(value or "").strip()
    if not value:
        return ""
    match = re.search(r"(\d{1,2})月(\d{1,2})日?", value)
    if match:
        return f"{int(match.group(1)):02d}-{int(match.group(2)):02d}"
    match = re.search(r"(\d{4})[-/](\d{1,2})[-/](\d{1,2})", value)
    if match:
        return f"{match.group(1)}-{int(match.group(2)):02d}-{int(match.group(3)):02d}"
    return value


def clean_excel_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%Y-%m-%d")
    if isinstance(value, (int, float)) and 25000 <= float(value) <= 60000:
        serial_date = datetime(1899, 12, 30) + timedelta(days=int(value))
        return serial_date.strftime("%Y-%m-%d")
    return re.sub(r"\s+", " ", str(value)).strip()


def find_plan_header_index(rows: list[list[str]]) -> int | None:
    for index, row in enumerate(rows[:20]):
        normalized = [clean_cell_text(cell) for cell in row]
        if "训练计划" in normalized or "训练任务" in normalized or "训练内容" in normalized:
            return index
    return None


def find_column(headers: list[str], names: list[str]) -> int | None:
    normalized = [clean_cell_text(header) for header in headers]
    for name in names:
        normalized_name = clean_cell_text(name)
        if normalized_name in normalized:
            return normalized.index(normalized_name)
    for index, header in enumerate(normalized):
        for name in names:
            if clean_cell_text(name) and clean_cell_text(name) in header:
                return index
    return None


def apply_optional_answer(record: dict[str, Any], name: str, value: str) -> None:
    value = str(value or "").strip()
    if not value:
        return
    if name == "星期":
        parsed_date = parse_record_date(value)
        if parsed_date:
            value = chinese_weekday(parsed_date)
    for answer in record["answers"]:
        if answer["name"] == name:
            answer["value"] = value
            answer["confidence"] = 0.75
            answer["reason"] = "从上传的训练计划 Excel 表格解析。"
            return


def parse_selected_modules(raw: str) -> list[str]:
    try:
        parsed = json.loads(raw) if raw.strip() else []
        if isinstance(parsed, list):
            return [str(item).strip().upper() for item in parsed if str(item).strip()]
    except json.JSONDecodeError:
        pass
    return []


def build_module_teacher_map(raw_map: str, module_a_teacher: str, module_b_teacher: str, module_c_teacher: str) -> dict[str, str]:
    mapping = {
        "A": module_a_teacher.strip(),
        "B": module_b_teacher.strip(),
        "C": module_c_teacher.strip(),
    }
    try:
        parsed = json.loads(raw_map) if raw_map.strip() else {}
        if isinstance(parsed, dict):
            mapping.update({str(key).strip().upper(): str(value).strip() for key, value in parsed.items()})
    except json.JSONDecodeError:
        pass
    return mapping


def apply_plan_defaults(
    records: list[dict[str, Any]],
    start_date: str,
    coach: str,
    writer: str,
    selected_modules: list[str],
    link_module_teacher: bool,
    module_teacher_map: dict[str, str],
) -> None:
    start = parse_start_date(start_date)
    for index, record in enumerate(records):
        parsed_module = get_record_answer(record, ["训练模块", "模块"])
        selected_module = selected_modules[index % len(selected_modules)] if selected_modules and not parsed_module else ""
        full_module = combine_module_name(selected_module, parsed_module)
        if start:
            current = start + timedelta(days=index)
            set_record_answer(record, "日期", f"{current.month}月{current.day}日", 0.8, "根据起始日期按顺序生成。")
            set_record_answer(record, "星期", chinese_weekday(current), 0.8, "根据起始日期按顺序生成。")
        if full_module:
            set_record_answer(record, "模块", full_module, 0.85, "来自页面模块选项和训练计划表格。")
            set_record_answer(record, "训练模块", full_module, 0.85, "来自页面模块选项和训练计划表格。")
            rewrite_generated_module_prefix(record, parsed_module, full_module)
        plan_title = first_plan_title([get_record_answer(record, ["训练任务", "训练计划"])])
        if plan_title:
            record["title"] = plan_title
        module_key = resolve_module_key(full_module or parsed_module or selected_module, selected_modules, module_teacher_map)
        linked_teacher = module_teacher_map.get(module_key, "") if link_module_teacher else ""
        effective_coach = linked_teacher or coach.strip()
        if effective_coach:
            reason = "来自模块-老师关联设置。" if linked_teacher else "来自页面默认设置。"
            set_record_answer(record, "授课老师", effective_coach, 0.85, reason)
            set_record_answer(record, "负责教练", effective_coach, 0.85, reason)
        if writer.strip():
            set_record_answer(record, "填写人", writer.strip(), 0.85, "来自页面默认设置。")


def parse_start_date(value: str) -> datetime | None:
    value = value.strip()
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            pass
    match = re.fullmatch(r"(\d{1,2})月(\d{1,2})日", value)
    if match:
        return datetime(datetime.now().year, int(match.group(1)), int(match.group(2)))
    return None


def chinese_weekday(value: datetime) -> str:
    return ["星期一", "星期二", "星期三", "星期四", "星期五", "星期六", "星期日"][value.weekday()]


def set_record_answer(record: dict[str, Any], name: str, value: str, confidence: float, reason: str) -> None:
    answers = record.setdefault("answers", [])
    for answer in answers:
        if answer.get("name") == name:
            answer["value"] = value
            answer["confidence"] = confidence
            answer["reason"] = reason
            return
    answers.append({"name": name, "value": value, "confidence": confidence, "reason": reason})


def get_record_answer(record: dict[str, Any], names: list[str]) -> str:
    return answer_list_value(record.get("answers", []), names)


def answer_list_value(answers: list[Any], names: list[str]) -> str:
    canonical_names = {canonical_answer_name(name) for name in names}
    for answer in answers:
        if (
            isinstance(answer, dict)
            and canonical_answer_name(str(answer.get("name", ""))) in canonical_names
            and str(answer.get("value", "")).strip()
        ):
            return str(answer.get("value", "")).strip()
    return ""


def combine_module_name(selected_module: str, parsed_module: str) -> str:
    selected_module = str(selected_module or "").strip()
    parsed_module = str(parsed_module or "").strip()
    if parsed_module:
        return parsed_module
    return selected_module


def resolve_module_key(module_value: str, selected_modules: list[str], module_teacher_map: dict[str, str]) -> str:
    normalized = clean_cell_text(module_value).upper()
    candidates = list(module_teacher_map.keys()) + selected_modules
    for candidate in candidates:
        key = clean_cell_text(candidate).upper()
        if key and (normalized == key or normalized.startswith(f"{key}-")):
            return key
    return normalized


def rewrite_generated_module_prefix(record: dict[str, Any], old_module: str, new_module: str) -> None:
    old_module = str(old_module or "").strip()
    new_module = str(new_module or "").strip()
    if not old_module or not new_module or clean_cell_text(old_module) == clean_cell_text(new_module):
        return
    for field_name in ["训练内容", "学习笔记"]:
        value = get_record_answer(record, [field_name])
        if not value:
            continue
        value = re.sub(rf"^{re.escape(old_module)}\s*专项训练", f"{new_module} 专项训练", value, count=1)
        set_record_answer(record, field_name, value, 0.85, "自动生成。")


def first_plan_title(plans: list[str]) -> str:
    for plan in plans:
        for line in str(plan or "").splitlines():
            cleaned = re.sub(r"^\s*\d+[.、]\s*", "", line).strip()
            if cleaned:
                return cleaned
    return ""


def parse_selected_records(raw_records: str) -> list[dict[str, Any]]:
    raw_records = raw_records.strip()
    if not raw_records:
        return []
    try:
        parsed = json.loads(raw_records)
        if not isinstance(parsed, list):
            raise ValueError("selected_records must be a list")
        records = []
        for index, record in enumerate(parsed, start=1):
            if not isinstance(record, dict):
                continue
            normalized_record = {
                "title": str(record.get("title") or f"training-log-{index}").strip(),
                "answers": normalize_answers(record.get("answers", [])),
            }
            if isinstance(record.get("source"), dict):
                normalized_record["source"] = record["source"]
            records.append(normalized_record)
        return records
    except (json.JSONDecodeError, ValueError, TypeError) as exc:
        raise HTTPException(status_code=400, detail=f"selected_records is invalid JSON: {exc}")


def save_api_config(config: dict[str, str], username: str) -> None:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    api_config_path_for_user(username).write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")


def parse_form_schema(raw_schema: str) -> list[dict[str, Any]]:
    raw_schema = raw_schema.strip()
    if not raw_schema:
        return []

    try:
        parsed = json.loads(raw_schema)
        if isinstance(parsed, dict):
            parsed = parsed.get("fields", [])
        if not isinstance(parsed, list):
            raise ValueError("schema must be a list")
        fields = []
        for item in parsed:
            if isinstance(item, str):
                fields.append({"name": item, "options": []})
            elif isinstance(item, dict) and item.get("name"):
                fields.append(
                    {
                        "name": str(item["name"]).strip(),
                        "options": [str(option).strip() for option in item.get("options", [])],
                        "description": str(item.get("description", "")).strip(),
                    }
                )
        return [field for field in fields if field["name"]]
    except (json.JSONDecodeError, ValueError, TypeError):
        return parse_lines_schema(raw_schema)


def parse_lines_schema(raw_schema: str) -> list[dict[str, Any]]:
    fields = []
    for line in raw_schema.splitlines():
        line = line.strip()
        if not line:
            continue
        if ":" in line:
            name, options_text = line.split(":", 1)
            options = [part.strip() for part in re.split(r"[,，|]", options_text) if part.strip()]
            fields.append({"name": name.strip(), "options": options})
        else:
            fields.append({"name": line, "options": []})
    return fields


async def save_upload(file: UploadFile) -> Path:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in {".pdf", ".docx", ".txt", ".md"}:
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT, and MD files are supported.")

    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    target = UPLOAD_DIR / f"{uuid.uuid4().hex}{suffix}"
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size cannot exceed 20MB.")
    target.write_bytes(content)
    return target


async def resolve_template_file(file: UploadFile | None) -> tuple[Path, str]:
    if file and file.filename:
        return await save_upload(file), file.filename
    if not DEFAULT_TEMPLATE_PATH.exists():
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
    return DEFAULT_TEMPLATE_PATH, "默认训练日志模板.docx"


def extract_text(path: Path, filename: str) -> str:
    suffix = path.suffix.lower() or Path(filename).suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        doc = Document(str(path))
        paragraphs = [paragraph.text for paragraph in doc.paragraphs]
        table_cells = [
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells)
    return path.read_text(encoding="utf-8", errors="ignore")


def extract_template_preview_from_bytes(content: bytes, suffix: str) -> tuple[str, list[dict[str, Any]]]:
    if suffix == ".pdf":
        reader = PdfReader(BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages), []
    if suffix == ".docx":
        doc = Document(BytesIO(content))
        text = extract_docx_text_from_document(doc)
        return text, extract_docx_template_fields_from_document(doc, text)
    return content.decode("utf-8", errors="ignore"), []


def extract_docx_text_from_document(doc: Document) -> str:
    paragraphs = [paragraph.text for paragraph in doc.paragraphs]
    table_cells = [
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    ]
    return "\n".join(paragraphs + table_cells)


def extract_docx_template_fields_from_document(doc: Document, document_text: str = "") -> list[dict[str, Any]]:
    names: list[str] = []
    horizontal_headers = {"日期", "星期", "时间", "模块", "训练计划", "负责教练"}
    for table in doc.tables:
        for row in table.rows:
            texts = [clean_cell_text(cell.text) for cell in row.cells]
            non_empty = [text for text in texts if text]
            if not non_empty:
                continue
            if len(set(non_empty) & horizontal_headers) >= 2:
                names.extend(non_empty)
                continue
            for index, text in enumerate(texts):
                if is_field_label(text) and (index + 1 < len(texts)) and not texts[index + 1]:
                    names.append(text)
            if is_field_label(texts[0]) and len(texts) > 1:
                names.append(texts[0])
    fallback = ["日期", "星期", "时间", "模块", "训练计划", "负责教练", "授课老师", "训练模块", "训练任务", "训练内容", "学习笔记", "总结", "填写人"]
    names.extend([name for name in fallback if name in document_text])
    return [{"name": name, "options": []} for name in dedupe(names)]


def extract_docx_template_fields(path: Path) -> list[dict[str, Any]]:
    doc = Document(str(path))
    names: list[str] = []
    horizontal_headers = {"日期", "星期", "时间", "模块", "训练计划", "负责教练"}
    for table in doc.tables:
        for row in table.rows:
            texts = [clean_cell_text(cell.text) for cell in row.cells]
            non_empty = [text for text in texts if text]
            if not non_empty:
                continue
            if len(set(non_empty) & horizontal_headers) >= 2:
                names.extend(non_empty)
                continue
            for index, text in enumerate(texts):
                if is_field_label(text) and (index + 1 < len(texts)) and not texts[index + 1]:
                    names.append(text)
            if is_field_label(texts[0]) and len(texts) > 1:
                names.append(texts[0])

    fallback = ["日期", "星期", "时间", "模块", "训练计划", "负责教练", "授课老师", "训练模块", "训练任务", "训练内容", "学习笔记", "总结", "填写人"]
    names.extend([name for name in fallback if name in extract_text(path, path.name)])
    return [{"name": name, "options": []} for name in dedupe(names)]


def is_field_label(text: str) -> bool:
    if not text or len(text) > 12:
        return False
    known = {"日期", "星期", "时间", "模块", "训练计划", "负责教练", "授课老师", "训练模块", "训练任务", "训练内容", "学习笔记", "总结", "填写人"}
    if text in known:
        return True
    return bool(re.fullmatch(r"[\u4e00-\u9fffA-Za-z]+", text))


def clean_cell_text(text: str) -> str:
    text = re.sub(r"\s+", "", text or "")
    text = re.sub(r"[/\\|:：]", "", text)
    return text.strip()


async def fetch_requirements_url(
    url: str,
    cookie: str = "",
    user_agent: str = "",
    referer: str = "",
) -> str:
    url = url.strip()
    if not url:
        return ""
    if not re.match(r"^https?://", url, flags=re.IGNORECASE):
        raise HTTPException(status_code=400, detail="Requirements URL must start with http:// or https://.")
    if re.match(r"^https?://(localhost|127\.0\.0\.1)(:\d+)?/", url, flags=re.IGNORECASE):
        raise HTTPException(
            status_code=400,
            detail="Docker 容器里不能用 localhost 读取宿主机网页。请把本机链接改成 http://host.docker.internal:端口/路径。",
        )

    headers = {
        "User-Agent": user_agent.strip()
        or "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/125 Safari/537.36",
    }
    if cookie.strip():
        headers["Cookie"] = cookie.strip()
    if referer.strip():
        headers["Referer"] = referer.strip()

    try:
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
            response = await client.get(url, headers=headers)
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=400, detail=f"Requirements URL could not be read: {exc}") from exc
    if response.status_code >= 400:
        raise HTTPException(status_code=400, detail=f"Requirements URL returned HTTP {response.status_code}.")
    content_type = response.headers.get("content-type", "")
    if "text" not in content_type and "html" not in content_type and "json" not in content_type:
        raise HTTPException(status_code=400, detail="Requirements URL must return readable text, HTML, or JSON.")
    return strip_html(response.text)


def strip_html(text: str) -> str:
    text = re.sub(r"(?is)<script.*?</script>|<style.*?</style>", " ", text)
    text = re.sub(r"(?s)<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


async def generate_answers(
    document_text: str,
    requirements_text: str,
    fields: list[dict[str, Any]],
    api_config: dict[str, str],
    custom_prompt: str = "",
    custom_skills: str = "",
) -> dict[str, Any]:
    api_format = api_config["format"]
    api_key = api_config["api_key"]
    base_url = api_config["base_url"]
    model = api_config["model"]

    if api_format == "openai" and not api_key:
        started = time.perf_counter()
        answers = heuristic_answers(document_text, requirements_text, fields)
        return {"used_ai": False, "answers": answers, "generation_stats": fallback_generation_stats(time.perf_counter() - started)}

    prompt = build_prompt(document_text, requirements_text, fields, custom_prompt, custom_skills)
    if api_format == "ollama":
        return await call_ollama(base_url, model, prompt, api_config["api_key"])
    return await call_openai_compatible(base_url, api_key, model, prompt)


async def generate_batch_records(
    document_text: str,
    requirements_text: str,
    fields: list[dict[str, Any]],
    api_config: dict[str, str],
    custom_prompt: str = "",
    custom_skills: str = "",
) -> dict[str, Any]:
    if not requirements_text.strip():
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")

    api_format = api_config["format"]
    api_key = api_config["api_key"]
    base_url = api_config["base_url"]
    model = api_config["model"]

    if api_format == "openai" and not api_key:
        started = time.perf_counter()
        records = heuristic_batch_records(document_text, requirements_text, fields)
        attach_stats_to_records(records, fallback_generation_stats(time.perf_counter() - started))
        return {"used_ai": False, "records": records}

    prompt = build_batch_prompt(document_text, requirements_text, fields, custom_prompt, custom_skills)
    if api_format == "ollama":
        raw_result = await call_ollama_raw_with_stats(base_url, model, prompt, api_config["api_key"])
    else:
        raw_result = await call_openai_compatible_raw_with_stats(base_url, api_key, model, prompt)
    records = parse_batch_json(raw_result["content"])
    attach_stats_to_records(records, raw_result["generation_stats"])
    return {"used_ai": True, "records": records}


async def generate_selected_records_one_by_one(
    document_text: str,
    requirements_text: str,
    selected_records: list[dict[str, Any]],
    fields: list[dict[str, Any]],
    api_config: dict[str, str],
    custom_prompt: str = "",
    custom_skills: str = "",
) -> dict[str, Any]:
    records = []
    for index, record in enumerate(selected_records, start=1):
        selected_requirements = build_selected_records_requirements([record], requirements_text)
        try:
            result = await generate_batch_records(
                document_text,
                selected_requirements[:MAX_REQUIREMENTS_CHARS],
                fields,
                api_config,
                custom_prompt,
                custom_skills,
            )
        except HTTPException as exc:
            detail = exc.detail if isinstance(exc.detail, str) else str(exc.detail)
            title = record.get("title") or f"第 {index} 条记录"
            raise HTTPException(status_code=exc.status_code, detail=f"{title} 生成失败：{detail}") from exc
        else:
            generated = result.get("records", [])
        if generated:
            records.append(generated[0])
    return {"used_ai": True, "records": records}


async def call_openai_compatible(
    base_url: str,
    api_key: str,
    model: str,
    prompt: str,
) -> dict[str, Any]:
    raw_result = await call_openai_compatible_raw_with_stats(base_url, api_key, model, prompt)
    answers = parse_ai_json(raw_result["content"])
    return {"used_ai": True, "answers": answers, "generation_stats": raw_result["generation_stats"]}


def fallback_record_from_selected(record: dict[str, Any], fields: list[dict[str, Any]], index: int, reason: str) -> dict[str, Any]:
    answers = record.get("answers", [])
    if not isinstance(answers, list):
        answers = []
    module = answer_list_value(answers, ["训练模块", "模块"])
    plan = answer_list_value(answers, ["训练任务", "训练计划"]) or str(record.get("title") or "训练任务")
    fallback = build_training_record(module, plan, fields, index)
    fallback["title"] = str(record.get("title") or fallback.get("title") or f"training-log-{index}")
    for answer in fallback.get("answers", []):
        answer["reason"] = f"本地兜底生成：{reason}"
    for existing in normalize_answers(answers):
        if existing.get("value"):
            set_record_answer(fallback, existing["name"], existing["value"], existing.get("confidence", 0.7), existing.get("reason", "自动生成。"))
    return fallback


async def call_openai_compatible_raw(base_url: str, api_key: str, model: str, prompt: str) -> str:
    result = await call_openai_compatible_raw_with_stats(base_url, api_key, model, prompt)
    return result["content"]


async def call_openai_compatible_raw_with_stats(base_url: str, api_key: str, model: str, prompt: str) -> dict[str, Any]:
    started = time.perf_counter()
    try:
        timeout = httpx.Timeout(300.0, connect=30.0, read=300.0, write=30.0, pool=30.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(
                f"{base_url}/chat/completions",
                headers={"Authorization": f"Bearer {api_key}"},
                json={
                    "model": model,
                    "messages": [
                        {
                            "role": "system",
                            "content": "请根据输入生成训练日志内容。",
                        },
                        {"role": "user", "content": prompt},
                    ],
                    "temperature": 0.1,
                },
            )
    except httpx.ConnectError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"无法连接 AI API：{base_url}。如果 API 在本机运行，Docker 内应使用 http://host.docker.internal:端口/v1，而不是 localhost。",
        ) from exc
    except httpx.TimeoutException as exc:
        raise HTTPException(
            status_code=504,
            detail=f"AI API 请求超时：正式生成内容较长，模型在 300 秒内没有返回。当前模型：{model}。",
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"AI API 请求失败：{type(exc).__name__}: {exc}") from exc

    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=describe_ai_http_error(response.status_code, response.text, base_url, model))
    payload = response.json()
    content = payload["choices"][0]["message"]["content"]
    usage = payload.get("usage") if isinstance(payload.get("usage"), dict) else {}
    return {
        "content": content,
        "generation_stats": build_generation_stats(
            "openai-compatible",
            model,
            prompt,
            content,
            usage,
            time.perf_counter() - started,
            "api" if usage else "estimated",
        ),
    }


async def list_openai_compatible_models(base_url: str, api_key: str, current_model: str) -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.get(
                f"{base_url}/models",
                headers={"Authorization": f"Bearer {api_key}"},
            )
    except httpx.ConnectError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"无法连接 AI API：{base_url}。如果 API 在本机运行，Docker 内应使用 http://host.docker.internal:端口/v1，而不是 localhost。",
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"获取模型请求失败：{exc}") from exc

    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=describe_ai_http_error(response.status_code, response.text, base_url, current_model))
    return extract_model_ids(response.json())


async def call_ollama(base_url: str, model: str, prompt: str, api_key: str = "") -> dict[str, Any]:
    raw_result = await call_ollama_raw_with_stats(base_url, model, prompt, api_key)
    answers = parse_ai_json(raw_result["content"])
    return {"used_ai": True, "answers": answers, "generation_stats": raw_result["generation_stats"]}


def ollama_headers(api_key: str = "") -> dict[str, str]:
    return {"Authorization": f"Bearer {api_key}"} if api_key else {}


async def call_ollama_raw(base_url: str, model: str, prompt: str, api_key: str = "") -> str:
    result = await call_ollama_raw_with_stats(base_url, model, prompt, api_key)
    return result["content"]


async def call_ollama_raw_with_stats(base_url: str, model: str, prompt: str, api_key: str = "") -> dict[str, Any]:
    started = time.perf_counter()
    try:
        timeout = httpx.Timeout(300.0, connect=30.0, read=300.0, write=30.0, pool=30.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(
                f"{base_url}/api/chat",
                headers=ollama_headers(api_key),
                json={
                    "model": model,
                    "stream": False,
                    "format": "json",
                    "messages": [
                        {
                            "role": "system",
                            "content": "请根据输入生成训练日志内容。",
                        },
                        {"role": "user", "content": prompt},
                    ],
                    "options": {"temperature": 0.1},
                },
            )
    except httpx.ConnectError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"无法连接 Ollama：{base_url}。如果 Ollama 在本机运行，Docker 内应使用 http://host.docker.internal:11434，而不是 localhost。",
        ) from exc
    except httpx.TimeoutException as exc:
        raise HTTPException(
            status_code=504,
            detail=f"Ollama 请求超时：正式生成内容较长，模型在 300 秒内没有返回。当前模型：{model}。",
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"Ollama 请求失败：{type(exc).__name__}: {exc}") from exc

    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=describe_ai_http_error(response.status_code, response.text, base_url, model))
    payload = response.json()
    content = payload.get("message", {}).get("content", "")
    usage = {
        "prompt_eval_count": payload.get("prompt_eval_count"),
        "eval_count": payload.get("eval_count"),
        "total_tokens": payload.get("total_tokens"),
    }
    duration_seconds = (
        (float(payload.get("total_duration") or 0) / 1_000_000_000)
        or float(payload.get("duration_seconds") or 0)
        or (time.perf_counter() - started)
    )
    has_usage = bool(usage.get("prompt_eval_count") or usage.get("eval_count") or usage.get("total_tokens"))
    token_source = str(payload.get("token_source") or ("api" if has_usage else "estimated"))
    return {
        "content": content,
        "generation_stats": build_generation_stats(
            "ollama",
            model,
            prompt,
            content,
            usage,
            duration_seconds,
            token_source,
        ),
    }


async def ensure_ollama_model(config: dict[str, str]) -> None:
    if config["format"] != "ollama":
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
    models = await list_ollama_models(config["base_url"], config["api_key"])
    if config["model"] not in models:
        raise HTTPException(
            status_code=400,
            detail=f"Ollama 已连接，但当前模型不存在：{config['model']}。可用模型：{', '.join(models[:20]) or '无'}",
        )


async def set_ollama_model_keep_alive(base_url: str, model: str, keep_alive: str | int, api_key: str = "") -> None:
    try:
        timeout = httpx.Timeout(300.0, connect=30.0, read=300.0, write=30.0, pool=30.0)
        async with httpx.AsyncClient(timeout=timeout) as client:
            response = await client.post(
                f"{base_url}/api/generate",
                headers=ollama_headers(api_key),
                json={
                    "model": model,
                    "prompt": "",
                    "stream": False,
                    "keep_alive": keep_alive,
                },
            )
    except httpx.ConnectError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"无法连接 Ollama：{base_url}。如果 Ollama 在本机运行，Docker 内应使用 http://host.docker.internal:11434，而不是 localhost。",
        ) from exc
    except httpx.TimeoutException as exc:
        raise HTTPException(status_code=504, detail="请求失败，请检查输入或 API 配置。") from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"Ollama 模型加载控制失败：{type(exc).__name__}: {exc}") from exc

    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=describe_ai_http_error(response.status_code, response.text, base_url, model))


async def list_ollama_models(base_url: str, api_key: str = "") -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.get(f"{base_url}/api/tags", headers=ollama_headers(api_key))
    except httpx.ConnectError as exc:
        raise HTTPException(
            status_code=502,
            detail=f"无法连接 Ollama：{base_url}。如果 Ollama 在本机运行，Docker 内应使用 http://host.docker.internal:11434，而不是 localhost。",
        ) from exc
    except httpx.HTTPError as exc:
        raise HTTPException(status_code=502, detail=f"Ollama 获取模型请求失败：{exc}") from exc

    if response.status_code >= 400:
        raise HTTPException(status_code=502, detail=f"Ollama 返回 HTTP {response.status_code}：{response.text[:500]}")
    return extract_model_ids(response.json())


async def detect_nvidia_gpus() -> list[dict[str, Any]]:
    def run_nvidia_smi() -> list[dict[str, Any]]:
        try:
            result = subprocess.run(
                ["nvidia-smi", "--query-gpu=name,index,memory.total", "--format=csv,noheader"],
                capture_output=True,
                text=True,
                timeout=10,
                check=False,
            )
        except (FileNotFoundError, subprocess.SubprocessError):
            return []
        if result.returncode != 0:
            return []
        gpus: list[dict[str, Any]] = []
        for line in result.stdout.splitlines():
            parts = [part.strip() for part in line.split(",")]
            if len(parts) < 2 or not parts[0]:
                continue
            gpus.append({
                "name": parts[0],
                "index": parts[1],
                "memory": parts[2] if len(parts) > 2 else "",
            })
        return gpus

    return await asyncio.to_thread(run_nvidia_smi)


async def read_ollama_processors(base_url: str) -> list[str]:
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            response = await client.get(f"{base_url}/api/ps")
    except httpx.HTTPError:
        return []
    if response.status_code >= 400:
        return []
    payload = response.json()
    processors: list[str] = []
    for item in payload.get("models", []):
        processor = str(item.get("processor") or "").strip()
        if processor and processor not in processors:
            processors.append(processor)
    return processors


def extract_model_ids(payload: Any) -> list[str]:
    if isinstance(payload, list):
        raw_items = payload
    elif isinstance(payload, dict):
        raw_items = payload.get("data") or payload.get("models") or payload.get("items") or []
    else:
        raw_items = []

    models: list[str] = []
    for item in raw_items:
        if isinstance(item, str):
            model_id = item
        elif isinstance(item, dict):
            model_id = str(item.get("id") or item.get("name") or item.get("model") or "").strip()
        else:
            model_id = ""
        if model_id:
            models.append(model_id)
    return sorted(dedupe_preserve_text(models), key=str.lower)


def dedupe_preserve_text(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        normalized = item.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result


def describe_ai_http_error(status_code: int, body: str, base_url: str, model: str) -> str:
    prefix = f"AI API 返回 HTTP {status_code}."
    if status_code in {401, 403}:
        hint = "请检查 API 地址、模型名称、密钥和兼容格式。"
    elif status_code == 404:
        hint = "请检查 API 地址、模型名称、密钥和兼容格式。"
    elif status_code == 429:
        hint = "请检查 API 地址、模型名称、密钥和兼容格式。"
    elif status_code >= 500:
        hint = "请检查 API 地址、模型名称、密钥和兼容格式。"
    else:
        hint = "请检查 API 地址、模型名称、密钥和兼容格式。"
    return f"{prefix}{hint} 返回内容：{body[:500]}"


def parse_ai_json(content: str) -> list[dict[str, Any]]:
    content = clean_json_content(content)
    try:
        parsed = json.loads(content)
        answers = parsed.get("answers", parsed)
        if not isinstance(answers, list):
            raise ValueError("answers must be a list")
        return normalize_answers(answers)
    except (json.JSONDecodeError, ValueError, TypeError) as exc:
        raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {exc}")


def parse_batch_json(content: str) -> list[dict[str, Any]]:
    content = clean_json_content(content)
    try:
        parsed = json.loads(content)
        records = parsed.get("records", parsed)
        if not isinstance(records, list):
            raise ValueError("records must be a list")
        normalized = []
        for index, record in enumerate(records, start=1):
            if not isinstance(record, dict):
                continue
            answers = record.get("answers", [])
            if not isinstance(answers, list):
                answers = []
            normalized.append(
                {
                    "title": str(record.get("title") or f"training-log-{index}").strip(),
                    "answers": normalize_answers(answers),
                }
            )
        return normalized
    except (json.JSONDecodeError, ValueError, TypeError) as exc:
        raise HTTPException(status_code=502, detail=f"AI returned invalid batch JSON: {exc}")


def clean_json_content(content: str) -> str:
    content = content.strip()
    if content.startswith("```"):
        content = re.sub(r"^```(?:json)?\s*|\s*```$", "", content, flags=re.IGNORECASE | re.DOTALL).strip()
    match = re.search(r"(\{.*\}|\[.*\])", content, flags=re.DOTALL)
    return match.group(1).strip() if match else content


def normalize_answers(answers: list[Any]) -> list[dict[str, Any]]:
    normalized = []
    for item in answers:
        if not isinstance(item, dict) or not item.get("name"):
            continue
        normalized.append(
            {
                "name": str(item.get("name", "")).strip(),
                "value": str(item.get("value", "")).strip(),
                "confidence": float(item.get("confidence", 0.0) or 0.0),
                "reason": str(item.get("reason", "")).strip(),
            }
        )
    return normalized


def build_prompt(
    document_text: str,
    requirements_text: str,
    fields: list[dict[str, Any]],
    custom_prompt: str = "",
    custom_skills: str = "",
) -> str:
    return json.dumps(
        {
            "task": (
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
            ),
            "output_format": {
                "answers": [
                    {
                        "name": "字段名",
                        "value": "瑕佸啓鍏ユā鏉跨殑鍐呭",
                        "confidence": 0.0,
                        "reason": "简短依据",
                    }
                ]
            },
            "fields": fields,
            "style_guide": "请根据输入生成训练日志内容。",
            "custom_prompt": custom_prompt.strip(),
            "custom_skills": parse_custom_skills(custom_skills),
            "monthly_requirements": requirements_text,
            "document_text": document_text,
        },
        ensure_ascii=False,
    )


def build_batch_prompt(
    document_text: str,
    requirements_text: str,
    fields: list[dict[str, Any]],
    custom_prompt: str = "",
    custom_skills: str = "",
) -> str:
    return json.dumps(
        {
            "task": (
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
                "请根据输入生成训练日志内容。",
            ),
            "output_format": {
                "records": [
                    {
                        "title": "训练日志",
                        "answers": [
                            {
                                "name": "字段名",
                                "value": "瑕佸啓鍏ユā鏉跨殑鍐呭",
                                "confidence": 0.0,
                                "reason": "简短依据",
                            }
                        ],
                    }
                ]
            },
            "fields": fields,
            "style_guide": "请根据输入生成训练日志内容。",
            "custom_prompt": custom_prompt.strip(),
            "custom_skills": parse_custom_skills(custom_skills),
            "monthly_requirements": requirements_text,
            "template_text": document_text,
        },
        ensure_ascii=False,
    )


def parse_custom_skills(raw: str) -> list[str]:
    seen = set()
    skills = []
    for line in raw.splitlines():
        item = line.strip()
        if not item or item in seen:
            continue
        seen.add(item)
        skills.append(item)
    return skills


def heuristic_answers(document_text: str, requirements_text: str, fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    combined = f"{document_text}\n{requirements_text}"
    normalized = combined.lower()
    answers = []
    for field in fields:
        value = ""
        confidence = 0.15
        explicit_value = extract_field_value(combined, field["name"])
        if explicit_value:
            value = explicit_value[:500]
            confidence = 0.45
            for option in field.get("options", []):
                if option.lower() in explicit_value.lower():
                    value = option
                    confidence = 0.65
                    break
        if not value:
            for option in field.get("options", []):
                if option.lower() in normalized:
                    value = option
                    confidence = 0.35
                    break
        answers.append(
            {
                "name": field["name"],
                "value": value,
                "confidence": confidence,
                "reason": "AI key is not configured; local fallback matching was used.",
            }
        )
    return answers


def heuristic_batch_records(document_text: str, requirements_text: str, fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    table_records = parse_training_plan_table(requirements_text, fields)
    if table_records:
        return table_records
    chunks = split_requirement_chunks(requirements_text)
    return [
        {
            "title": f"training-log-{index}",
            "answers": heuristic_answers(document_text, chunk, fields),
        }
        for index, chunk in enumerate(chunks, start=1)
    ]


def parse_training_plan_table(text: str, fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    records = []
    for line in text.splitlines():
        raw = line.strip()
        if not raw:
            continue
        if "训练计划" in raw and ("模块" in raw or "日期" in raw):
            continue
        parts = [part.strip() for part in re.split(r"\t+|\s{2,}|[|]", raw) if part.strip()]
        module = ""
        plan = ""
        if len(parts) >= 2:
            module = parts[0]
            plan = parts[-1]
        elif re.match(r"^(Linux|Network|Windows)\b", raw, flags=re.IGNORECASE):
            match = re.match(r"^(Linux|Network|Windows)\s+(.+)$", raw, flags=re.IGNORECASE)
            if match:
                module = match.group(1)
                plan = match.group(2).strip()
        if not plan or len(plan) < 4:
            continue
        records.append(build_training_record(module, plan, fields, len(records) + 1))
    return records[:60]


def build_training_record(module: str, plan: str, fields: list[dict[str, Any]], index: int) -> dict[str, Any]:
    title = build_plan_title(module, plan, index)
    plan_lines = split_plan_lines(plan)
    content = build_numbered_training_content(title, plan_lines)
    notes = build_learning_notes(title, plan_lines)
    summary = build_summary_text(title, plan_lines)
    values = {
        "标题": title,
        "模块": module,
        "训练模块": module,
        "训练计划": plan,
        "训练任务": plan,
        "训练内容": content,
        "学习笔记": notes,
        "总结": summary,
    }
    answers = []
    for field in fields:
        name = field["name"]
        value = values.get(name, "")
        answers.append(
            {
                "name": name,
                "value": value,
                "confidence": 0.6 if value else 0.15,
                "reason": "从粘贴的训练计划表格本地解析。" if value else "本地解析未找到对应值。",
            }
        )
    return {"title": f"{index:02d}-{module or '训练'}-{plan[:28]}", "answers": answers}


def build_plan_title(module: str, plan: str, index: int) -> str:
    first_line = split_plan_lines(plan)[0] if split_plan_lines(plan) else plan
    prefix = module or "训练"
    topic = extract_topic(first_line)
    return f"training-log-{index}"


def split_plan_lines(plan: str) -> list[str]:
    lines = []
    for line in str(plan or "").splitlines():
        cleaned = re.sub(r"^\s*\d+[.、]\s*", "", line).strip()
        if cleaned:
            lines.append(cleaned)
    return lines or [str(plan or "").strip()]


def extract_topic(text: str) -> str:
    text = re.sub(r"\s+", " ", str(text or "")).strip()
    text = re.sub(r"^(完成|学习|掌握|理解|配置|训练|练习|熟悉|了解|实现)\s*", "", text)
    text = re.split(r"[。；;，,]", text, maxsplit=1)[0].strip()
    return text or str(text or "").strip()


def join_topics(plan_lines: list[str], limit: int = 3) -> str:
    topics = []
    seen = set()
    for line in plan_lines:
        topic = extract_topic(line)
        normalized = clean_cell_text(topic)
        if topic and normalized not in seen:
            seen.add(normalized)
            topics.append(topic)
    topics = topics[:limit]
    if not topics:
        return "本次训练任务"
    if len(topics) == 1:
        return topics[0]
    return "、".join(topics[:-1]) + "和" + topics[-1]


def build_numbered_training_content(title: str, plan_lines: list[str]) -> str:
    focus = join_topics(plan_lines)
    return (
        f"{title}\n"
        f"本次训练围绕{focus}开展专项学习和实操，记录关键概念、配置步骤、验证结果和排错过程。"
    )


def build_learning_notes(title: str, plan_lines: list[str]) -> str:
    focus = join_topics(plan_lines)
    details = concrete_learning_points(focus)
    sections = [
        f"1. 明确训练目标\n本次学习围绕{focus}展开，重点记录{details['goal']}。",
        f"2. 梳理核心组成\n训练中先区分{details['components']}，避免把配置步骤混在一起。",
        f"3. 完成配置实施\n实操时按顺序检查{details['implementation']}，并记录关键参数和位置。",
        f"4. 做好结果验证\n验证阶段重点查看{details['verification']}，异常时先确认基础连通和服务状态。",
        f"5. 记录排错思路\n排错时按照{details['troubleshooting']}的顺序处理，每次只调整一个变量。",
        f"6. 整理训练收获\n通过本次训练，对{details['takeaway']}有了更具体的理解。",
    ]
    return f"{title}\n" + "\n\n".join(sections)


def concrete_learning_points(focus: str) -> dict[str, str]:
    text = clean_cell_text(focus).upper()
    return {
        "goal": f"{focus}涉及的配置对象、关键参数、验证指标和排错依据",
        "components": "配置对象、前置条件、执行步骤、验证结果和日志信息",
        "implementation": "环境参数、账号或接口信息、服务状态、策略条件和关键命令",
        "verification": "服务状态、连通性、日志输出、权限结果和最终业务效果",
        "troubleshooting": "先检查基础环境，再检查配置参数、权限策略、端口连通和日志报错",
        "takeaway": f"{focus}从配置到验证再到排错的完整处理思路",
    }


def build_summary_text(title: str, plan_lines: list[str]) -> str:
    focus = join_topics(plan_lines)
    return (
        f"训练目标达成情况：本次训练围绕{focus}完成学习和实操，基本掌握任务的核心概念、配置流程和验证方法。\n\n"
        "不足之处：部分细节仍需要继续加强，例如关键参数的适用条件、配置顺序对结果的影响，以及故障出现时如何更快定位到具体环节。\n\n"
        "后续改进：继续补充操作截图、命令输出和排错记录，形成更完整的复盘材料。"
    )


def split_requirement_chunks(text: str) -> list[str]:
    lines = [line.strip() for line in re.split(r"[\r\n]+", text) if line.strip()]
    date_like = [line for line in lines if re.search(r"\d{1,2}[月/-]\d{1,2}|\d{4}[/-]\d{1,2}[/-]\d{1,2}|周[一二三四五六日天]", line)]
    if len(date_like) >= 2:
        return date_like[:31]
    inline_parts = [
        part.strip()
        for part in re.split(r"(?=\d{1,2}月\d{1,2}日?[：:])|(?=\d{4}[/-]\d{1,2}[/-]\d{1,2}[：:])|(?=周[一二三四五六日天][：:])", text)
        if part.strip()
    ]
    if len(inline_parts) >= 2:
        return inline_parts[:31]
    return [text[:MAX_REQUIREMENTS_CHARS]]


def extract_field_value(document_text: str, field_name: str) -> str:
    pattern = rf"^\s*{re.escape(field_name)}\s*[:：]\s*(.+?)\s*$"
    match = re.search(pattern, document_text, flags=re.MULTILINE)
    return match.group(1).strip() if match else ""


def fill_docx_template(template_path: Path, answers: list[dict[str, Any]], output_title: str = "") -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(template_path))
    answer_map = build_answer_map(answers)
    fill_document_paragraphs(doc, answer_map)

    for table in doc.tables:
        fill_table_with_answers(table, answer_map)

    suffix = f"-{safe_filename(output_title)}" if output_title else ""
    output_path = OUTPUT_DIR / f"{uuid.uuid4().hex}{suffix}.docx"
    doc.save(str(output_path))
    return output_path


def fill_docx_batch(template_path: Path, records: list[dict[str, Any]]) -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_paths = []
    for index, record in enumerate(records, start=1):
        output_path = fill_docx_template(template_path, record.get("answers", []), record.get("title", ""))
        safe_title = safe_filename(record.get("title") or f"training-log-{index}")
        renamed_path = OUTPUT_DIR / f"{uuid.uuid4().hex}-{safe_title}.docx"
        output_path.replace(renamed_path)
        docx_paths.append(renamed_path)

    zip_path = OUTPUT_DIR / f"{uuid.uuid4().hex}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for index, path in enumerate(docx_paths, start=1):
            archive.write(path, arcname=f"{index:02d}-{path.name.split('-', 1)[-1]}")
    return zip_path


def package_batch_outputs(template_path: Path, records: list[dict[str, Any]], output_format: str, package_mode: str) -> Path:
    if not records:
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
    normalized_format = (output_format or "docx").strip().lower()
    if normalized_format not in {"docx", "word", "pdf"}:
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    package_title = package_title_for_records(records)
    root_folder = safe_filename(package_title)
    used_names: set[str] = set()
    generated_paths: list[tuple[Path, str]] = []

    for index, record in enumerate(records, start=1):
        output_title = record_file_title(record, index)
        output_path = fill_docx_template(template_path, record.get("answers", []), output_title)
        if normalized_format == "pdf":
            output_path = convert_docx_to_pdf(output_path)
        file_name = unique_package_filename(f"{safe_filename(output_title)}{output_path.suffix}", used_names)
        generated_paths.append((output_path, file_name))

    if normalize_package_mode(package_mode) == "folder":
        folder_path = unique_output_folder(root_folder)
        folder_path.mkdir(parents=True, exist_ok=False)
        for path, file_name in generated_paths:
            shutil.copy2(path, folder_path / file_name)
        return folder_path

    zip_path = OUTPUT_DIR / f"{uuid.uuid4().hex}-{root_folder}.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path, file_name in generated_paths:
            arcname = f"{root_folder}/{file_name}"
            archive.write(path, arcname=arcname)
    return zip_path


def unique_package_filename(file_name: str, used_names: set[str]) -> str:
    candidate = file_name
    stem = Path(file_name).stem
    suffix = Path(file_name).suffix
    counter = 2
    while candidate in used_names:
        candidate = f"{stem}-{counter}{suffix}"
        counter += 1
    used_names.add(candidate)
    return candidate


def fill_docx_records(template_path: Path, records: list[dict[str, Any]]) -> Path:
    if not records:
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    base_doc: Any | None = None
    for index, record in enumerate(records):
        doc = Document(str(template_path))
        answers = record.get("answers", [])
        if not isinstance(answers, list):
            answers = []
        answer_map = build_answer_map(answers)
        fill_document_paragraphs(doc, answer_map)
        for table in doc.tables:
            fill_table_with_answers(table, answer_map)
        if base_doc is None:
            base_doc = doc
            continue
        base_doc.add_page_break()
        append_doc_body(base_doc, doc)

    output_stem = output_title_for_records(records)
    suffix = f"-{safe_filename(output_stem)}" if output_stem else ""
    output_path = OUTPUT_DIR / f"{uuid.uuid4().hex}{suffix}.docx"
    assert base_doc is not None
    base_doc.save(str(output_path))
    return output_path


ANSWER_ALIASES = {
    "训练模块": ["模块", "课程模块"],
    "训练计划": ["训练任务", "任务", "计划", "训练项目"],
    "训练任务": ["训练计划", "任务", "计划", "训练项目"],
    "训练内容": ["训练过程", "训练步骤", "实训内容", "实训过程", "实践内容"],
    "学习笔记": ["学习记录", "学习心得", "学习内容", "笔记", "实训笔记", "训练笔记"],
    "总结": ["训练总结", "学习总结", "实训总结", "小结", "训练目标达成情况", "目标达成情况"],
    "日期": ["时间", "训练日期"],
    "星期": ["周次", "周几"],
    "填写人": ["填表人", "记录人"],
}

CANONICAL_ANSWER_NAMES = {
    clean_cell_text(alias): canonical
    for canonical, aliases in ANSWER_ALIASES.items()
    for alias in [canonical, *aliases]
}


def canonical_answer_name(name: str) -> str:
    cleaned = clean_cell_text(name)
    return CANONICAL_ANSWER_NAMES.get(cleaned, cleaned)


def build_answer_map(answers: list[dict[str, Any]]) -> dict[str, str]:
    answer_map: dict[str, str] = {}
    summary_parts: list[str] = []
    for item in answers:
        if not isinstance(item, dict) or not item.get("name"):
            continue
        raw_name = clean_cell_text(str(item.get("name", "")))
        name = canonical_answer_name(raw_name)
        value = normalize_output_value(item.get("value", ""))
        if not value:
            continue
        if raw_name in {"训练目标达成情况", "目标达成情况", "不足之处", "改进建议"}:
            summary_parts.append(f"{raw_name}：{value}")
        answer_map.setdefault(raw_name, value)
        answer_map.setdefault(name, value)

    if "总结" not in answer_map and summary_parts:
        answer_map["总结"] = "\n\n".join(summary_parts)
    mirror_answer_value(answer_map, "模块", "训练模块")
    mirror_answer_value(answer_map, "训练计划", "训练任务")
    mirror_answer_value(answer_map, "授课老师", "负责教练")

    module = answer_map.get("训练模块") or answer_map.get("模块") or ""
    plan = answer_map.get("训练任务") or answer_map.get("训练计划") or answer_map.get("训练内容") or ""
    if plan:
        fallback = fallback_content_map(module, plan)
        for key in ["训练内容", "学习笔记", "总结"]:
            if not answer_map.get(key):
                answer_map[key] = fallback[key]
    return answer_map


def enrich_answers_for_fields(answers: list[Any], fields: list[dict[str, Any]]) -> list[dict[str, Any]]:
    normalized = normalize_answers(answers if isinstance(answers, list) else [])
    answer_map = build_answer_map(normalized)
    enriched: list[dict[str, Any]] = []
    seen: set[str] = set()

    for answer in normalized:
        name = clean_cell_text(answer.get("name", ""))
        canonical = canonical_answer_name(name)
        value = answer.get("value", "") or answer_map.get(name, "") or answer_map.get(canonical, "")
        if value:
            answer = {**answer, "value": value}
        enriched.append(answer)
        seen.add(name)

    for field in fields:
        name = clean_cell_text(str(field.get("name", "")))
        if not name or name in seen:
            continue
        value = answer_map.get(name) or answer_map.get(canonical_answer_name(name)) or ""
        if not value:
            continue
        enriched.append({
            "name": str(field.get("name", "")).strip(),
            "value": value,
            "confidence": 0.55,
            "reason": "简短依据",
        })
        seen.add(name)
    return enriched


def fallback_content_map(module: str, plan: str) -> dict[str, str]:
    title = build_plan_title(module, plan, 1)
    plan_lines = split_plan_lines(plan)
    return {
        "训练内容": build_numbered_training_content(title, plan_lines),
        "学习笔记": build_learning_notes(title, plan_lines),
        "总结": build_summary_text(title, plan_lines),
    }


def mirror_answer_value(answer_map: dict[str, str], left: str, right: str) -> None:
    if answer_map.get(left) and not answer_map.get(right):
        answer_map[right] = answer_map[left]
    if answer_map.get(right) and not answer_map.get(left):
        answer_map[left] = answer_map[right]


def append_doc_body(target_doc: Document, source_doc: Document) -> None:
    from copy import deepcopy

    target_body = target_doc.element.body
    for element in source_doc.element.body:
        if element.tag.endswith("}sectPr"):
            continue
        target_body.append(deepcopy(element))


def convert_output_format(docx_path: Path, output_format: str) -> Path:
    normalized = (output_format or "docx").strip().lower()
    if normalized in {"docx", "word"}:
        return docx_path
    if normalized != "pdf":
        raise HTTPException(status_code=400, detail="请求失败，请检查输入或 API 配置。")
    return convert_docx_to_pdf(docx_path)


def convert_docx_to_pdf(docx_path: Path) -> Path:
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice:
        raise HTTPException(status_code=500, detail="请求失败，请检查输入或 API 配置。")
    result = subprocess.run(
        [
            soffice,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(OUTPUT_DIR),
            str(docx_path),
        ],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=90,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    if result.returncode != 0 or not pdf_path.exists():
        detail = (result.stderr or result.stdout or "LibreOffice 没有返回可用错误信息。").strip()
        raise HTTPException(status_code=500, detail=f"PDF 导出失败：{detail[:500]}")
    return pdf_path


def title_from_answers(answers: list[dict[str, Any]], fallback: str = "") -> str:
    title = answer_list_value(answers, ["训练任务", "训练计划"])
    return first_plan_title([title]) or Path(fallback).stem


def output_title_for_records(records: list[dict[str, Any]]) -> str:
    if not records:
        return "training-log"
    first_title = first_plan_title([str(records[0].get("title", ""))])
    if len(records) == 1:
        return first_title or title_from_answers(records[0].get("answers", []), "training-log")
    return f"{first_title or '训练日志'}-等{len(records)}篇"


def package_title_for_records(records: list[dict[str, Any]]) -> str:
    dates = [parse_record_date(record_date_value(record)) for record in records]
    if dates and all(dates):
        typed_dates = [date for date in dates if date is not None]
        if len(typed_dates) == 1:
            return f"{format_record_date(typed_dates[0])}训练日志"
        if are_continuous_dates(typed_dates):
            return f"{format_record_date(typed_dates[0])}-{format_record_date(typed_dates[-1])}训练日志"
        return f"{format_record_date(typed_dates[0])}等{len(typed_dates)}天训练日志"
    return f"{output_title_for_records(records)}训练日志"


def record_file_title(record: dict[str, Any], index: int) -> str:
    parsed_date = parse_record_date(record_date_value(record))
    if parsed_date:
        return f"{format_record_date(parsed_date)}训练日志"
    fallback = first_plan_title([str(record.get("title", ""))]) or title_from_answers(record.get("answers", []), f"{index:02d}训练日志")
    return f"{fallback}训练日志" if "训练日志" not in fallback else fallback


def record_date_value(record: dict[str, Any]) -> str:
    return answer_list_value(record.get("answers", []), ["日期", "训练日期", "时间"])


def parse_record_date(value: str) -> datetime | None:
    value = str(value or "").strip()
    if not value:
        return None
    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return datetime.strptime(value, fmt)
        except ValueError:
            pass
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日?", value)
    if match:
        return datetime(int(match.group(1)), int(match.group(2)), int(match.group(3)))
    match = re.search(r"(\d{1,2})月(\d{1,2})日?", value)
    if match:
        return datetime(datetime.now().year, int(match.group(1)), int(match.group(2)))
    return None


def are_continuous_dates(dates: list[datetime]) -> bool:
    if len(dates) < 2:
        return True
    return all((dates[index] - dates[index - 1]).days == 1 for index in range(1, len(dates)))


def format_record_date(value: datetime) -> str:
    return value.strftime("%Y-%m-%d")


def download_display_name(path: Path) -> str:
    stem = path.stem
    match = re.fullmatch(r"[a-f0-9]{32}-(.+)", stem)
    display_stem = match.group(1) if match else "训练日志"
    return f"{display_stem}{path.suffix}"


def safe_filename(value: str) -> str:
    value = re.sub(r"[^\w\u4e00-\u9fff.-]+", "-", value.strip())
    return value.strip("-")[:60] or "training-log"


def fill_document_paragraphs(doc: Document, answer_map: dict[str, str]) -> None:
    writer = answer_map.get("填写人", "").strip()
    if not writer:
        return
    for paragraph in doc.paragraphs:
        fill_writer_title_paragraph(paragraph, writer)


def fill_writer_title_paragraph(paragraph: Any, writer: str) -> None:
    text = paragraph.text
    if "选手日志" not in text:
        return
    updated = re.sub(r"(选手日志[：:])[^：:]*([：:])", lambda match: f"{match.group(1)}{writer}{match.group(2)}", text)
    if updated != text:
        set_paragraph_text(paragraph, updated)


def set_paragraph_text(paragraph: Any, value: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(value)
        return
    paragraph.runs[0].text = value
    for run in paragraph.runs[1:]:
        run.text = ""


def fill_table_with_answers(table: Any, answer_map: dict[str, str]) -> None:
    used_cells: set[int] = set()
    for row in table.rows:
        cells = row.cells
        texts = [clean_cell_text(cell.text) for cell in cells]
        fill_label_value_row(cells, texts, answer_map, used_cells)
        fill_section_row(cells, texts, answer_map, used_cells)
    if not used_cells:
        fill_horizontal_table(table, answer_map, used_cells)


def fill_horizontal_table(table: Any, answer_map: dict[str, str], used_cells: set[int]) -> bool:
    if not table.rows:
        return False
    headers = [clean_cell_text(cell.text) for cell in table.rows[0].cells]
    matched_indexes = [(index, header) for index, header in enumerate(headers) if header in answer_map]
    if len(matched_indexes) < 2:
        return False
    if any(index + 1 < len(headers) and not headers[index + 1] for index, _ in matched_indexes):
        return False

    target_row = None
    for row in table.rows[1:]:
        if any(not clean_cell_text(cell.text) for cell in row.cells):
            target_row = row
            break
    if target_row is None:
        target_row = table.add_row()

    for index, header in matched_indexes:
        if index < len(target_row.cells) and id(target_row.cells[index]) not in used_cells:
            set_cell_text(target_row.cells[index], answer_map[header])
            used_cells.add(id(target_row.cells[index]))
    return True


def fill_label_value_row(cells: list[_Cell], texts: list[str], answer_map: dict[str, str], used_cells: set[int]) -> None:
    for index, text in enumerate(texts[:-1]):
        value = answer_map.get(text)
        target = cells[index + 1]
        if value and not texts[index + 1] and id(target) not in used_cells:
            set_cell_text(target, value)
            used_cells.add(id(target))


def fill_section_row(cells: list[_Cell], texts: list[str], answer_map: dict[str, str], used_cells: set[int]) -> None:
    if not texts:
        return
    label = texts[0]
    value = answer_map.get(label)
    if not value or len(cells) < 2:
        return
    target = cells[1]
    if id(target) in used_cells:
        return
    set_cell_text(target, value)
    used_cells.add(id(target))


def set_cell_text(cell: _Cell, value: str) -> None:
    cell.text = value


def normalize_output_value(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s*00:00:00\s*$", "", text)
    return text


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    result = []
    for item in items:
        normalized = clean_cell_text(item)
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result
